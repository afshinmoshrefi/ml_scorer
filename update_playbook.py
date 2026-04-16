"""Add Section 11 (Backtest Results) to the Stock Strategy Playbook V2 document."""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document("docs/TradeWave_Stock_Strategy_Playbook_V2.docx")

# Remove duplicate style definitions (causes python-docx lookup failure)
_seen_styles = set()
_to_remove = []
for child in doc.styles.element:
    sid = child.get(qn("w:styleId"))
    if sid and sid in _seen_styles:
        _to_remove.append(child)
    elif sid:
        _seen_styles.add(sid)
for elem in _to_remove:
    doc.styles.element.remove(elem)


def add_heading(text, level=1):
    """Add heading using XML-level style to avoid python-docx duplicate style bug."""
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pStyle")):
        pPr.remove(existing)
    pPr.insert(0, pPr.makeelement(qn("w:pStyle"), {qn("w:val"): f"Heading{level}"}))
    return p

# ============================================================
# Section 11: Backtest Results
# ============================================================

doc.add_page_break()
add_heading("11. Backtest Results", level=1)

doc.add_paragraph(
    "This section documents the complete results of running all 112 strategy variants against "
    "the 8-year walk-forward validation dataset (2018-2025). The backtester was implemented per "
    "the specification in Section 10 and executed on 2026-03-20. All results use the 10-30 day "
    "tier with SR + MFE ensemble models trained on S&P 500 stocks."
)

# --- 11.1 ---
add_heading("11.1 Execution Summary", level=2)
doc.add_paragraph(
    "Input data: 8,039,032 long opportunities across 8 validation years (2018-2025), "
    "filtered from 11M walk-forward predictions (VIX > 35 removed, longs only). "
    "475 S&P 500 symbols, 1,949 unique trading dates with candidates out of 2,011 total trading days."
)
doc.add_paragraph(
    "Starting capital: $100,000. Slippage: 0.2% round-trip (0.1% entry + 0.1% exit). "
    "Cash reserve: 10% minimum. Hard stop: 10% of equity per position. "
    "Drawdown halt: 15% triggers 20-day trading pause."
)
doc.add_paragraph(
    "Runtime: 12 minutes on 12 parallel workers (24-core machine). "
    "Output: 76,709 total trades, 225,232 daily equity records across all 112 strategies."
)
doc.add_paragraph(
    "Filters applied: Direction = long only (playbook Section 4.3). "
    "VIX > 35 filter pre-applied in training data. "
    "Liquidity gate (volume, spread, market cap) skipped for backtest -- all S&P 500 stocks pass trivially. "
    "Earnings exclusion skipped -- historical earnings calendar data not available for backtest period."
)

# --- 11.2 ---
add_heading("11.2 Aggregate Results", level=2)
doc.add_paragraph(
    "101 of 112 strategies (90%) were profitable in all 8 years. "
    "Every single strategy was profitable in at least 7 of 8 years. "
    "107 of 112 strategies achieved portfolio Sharpe ratio above 2.0. "
    "39 strategies achieved Sharpe above 3.0. "
    "9 strategies kept maximum drawdown below 15%. "
    "25 strategies kept maximum drawdown below 20%."
)
doc.add_paragraph(
    "The 11 strategies that had a losing year all lost money in 2019. "
    "2019 is a structural weakness: the walk-forward model trained on 2000-2018 produced "
    "compressed predictions for 2019, resulting in only 722 ML_85+ opportunities "
    "(vs 240K-320K in other years). "
    "Strategies using T70 threshold had sufficient 2019 candidates; T85/T90 strategies "
    "in some configurations had too few trades to avoid individual losing positions "
    "dragging the year negative."
)

# --- 11.3 ---
add_heading("11.3 Dimension Analysis", level=2)
doc.add_paragraph(
    "Each of the six strategy dimensions was analyzed in isolation to identify which "
    "parameter values produce the best risk-adjusted results."
)

# EXIT RULES
add_heading("Exit Rules", level=3)
t = doc.add_table(rows=5, cols=6)
# t.style applied after creation
for i, h in enumerate(["Exit", "Avg Sharpe", "Avg Max DD", "Avg WR", "Count", "Key Finding"]):
    t.rows[0].cells[i].text = h
rows = [
    ["EP (3% Trail)", "3.95", "10.2%", "58.1%", "4",
     "Best risk-adjusted. Caps loss at 3% below HWM. Low WR but extreme asymmetry."],
    ["EM (MFE Trail)", "3.00", "26.1%", "85.3%", "66",
     "Best win rate. Locks in gains once predicted return reached."],
    ["ET (Time+Trail)", "2.73", "22.8%", "72.7%", "27",
     "Middle ground. 60% time limit exits dead positions."],
    ["EH (Hold)", "2.03", "35.9%", "76.9%", "15",
     "Worst. No loss management. Baseline control only."],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "EP is the standout finding. By exiting any position that pulls back 3% from its high "
    "water mark, EP creates an asymmetric payoff profile: many small losses (capped at ~3.2%) "
    "and fewer but much larger wins (avg 9%, with outliers above 100%). "
    "This produces the lowest drawdowns and highest Sharpe ratios despite a sub-60% win rate. "
    "The EP strategies never trigger the 15% drawdown halt."
)

# SIZING
add_heading("Sizing Methods", level=3)
t = doc.add_table(rows=7, cols=5)
# t.style applied after creation
for i, h in enumerate(["Sizing", "Avg Sharpe", "Avg Max DD", "Avg WR", "Count"]):
    t.rows[0].cells[i].text = h
rows = [
    ["SK (Quarter Kelly)", "3.11", "17.4%", "80.6%", "24"],
    ["SA (Adaptive Kelly)", "3.02", "23.1%", "79.4%", "24"],
    ["SH (Half Kelly)", "2.89", "28.8%", "81.2%", "46"],
    ["SV (Vol-Inverse)", "2.71", "19.4%", "84.5%", "1"],
    ["SC (Confidence)", "2.52", "35.4%", "83.4%", "1"],
    ["SF (Equal Weight)", "2.06", "35.8%", "77.4%", "16"],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "Quarter Kelly (SK) dominates. It allocates 25% of the full Kelly fraction, which "
    "for typical ML_85 candidates (W=0.82, R=1.2) gives ~17% of equity per position. "
    "This is conservative enough to limit drawdowns while capturing meaningful compounding. "
    "Half Kelly (SH) compounds faster but with 65% higher drawdowns. "
    "Equal Weight (SF) has the worst Sharpe because it ignores model confidence."
)

# RANKING
add_heading("Ranking Methods", level=3)
t = doc.add_table(rows=7, cols=5)
# t.style applied after creation
for i, h in enumerate(["Ranking", "Avg Sharpe", "Avg Max DD", "Avg WR", "Count"]):
    t.rows[0].cells[i].text = h
rows = [
    ["CR (Return-Weighted)", "3.20", "24.8%", "83.3%", "12"],
    ["CW (WP-Weighted)", "3.05", "23.6%", "81.0%", "25"],
    ["MG (MFE Gap)", "2.72", "29.6%", "84.0%", "6"],
    ["WP (Win Probability)", "2.72", "22.9%", "77.9%", "39"],
    ["PR (Predicted Return)", "2.71", "31.2%", "80.6%", "20"],
    ["MS (ML Score)", "2.67", "33.9%", "80.0%", "10"],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "Composite rankings (CR and CW) outperform single-metric rankings. "
    "CR (0.30*WP + 0.50*PR + 0.20*MG) produces the highest Sharpe by balancing win "
    "probability with expected magnitude and MFE runway."
)

# CONCENTRATION
add_heading("Sector Concentration", level=3)
t = doc.add_table(rows=4, cols=4)
# t.style applied after creation
for i, h in enumerate(["Rule", "Avg Sharpe", "Avg Max DD", "Count"]):
    t.rows[0].cells[i].text = h
rows = [
    ["C1 (Max 1/sector)", "3.13", "17.6%", "18"],
    ["C2 (Max 2/sector)", "2.92", "25.5%", "69"],
    ["CN (No limit)", "2.42", "33.8%", "25"],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "Sector diversification is the single largest driver of drawdown reduction. "
    "C1 cuts average drawdown nearly in half vs no limits (17.6% vs 33.8%). "
    "Correlated sector exposure is the primary source of portfolio risk."
)

# THRESHOLDS
add_heading("ML Score Thresholds", level=3)
t = doc.add_table(rows=5, cols=5)
# t.style applied after creation
for i, h in enumerate(["Threshold", "Avg Sharpe", "Avg Max DD", "Avg WR", "Count"]):
    t.rows[0].cells[i].text = h
rows = [
    ["T80", "3.00", "29.4%", "82.1%", "14"],
    ["T90", "2.92", "21.2%", "77.4%", "14"],
    ["T70", "2.89", "28.4%", "82.0%", "13"],
    ["T85", "2.78", "26.0%", "80.0%", "71"],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "Threshold differences are small relative to other dimensions. T90 has the lowest "
    "drawdown (fewer but higher-conviction trades) while T80 has the highest Sharpe "
    "(more trades, better compounding)."
)

# --- 11.4 ---
add_heading("11.4 Top Strategy Profiles", level=2)

# STRATEGY 23
add_heading("Strategy #23 -- The Safest Cash Machine", level=3)
doc.add_paragraph(
    "Configuration: WP ranking, T90 threshold, EP exit (3% trail), "
    "Quarter Kelly sizing, 3 max positions, C1 (1 per sector)."
)
t = doc.add_table(rows=2, cols=8)
# t.style applied after creation
for i, h in enumerate(["Sharpe", "Max DD", "WR", "Trades", "Ann Ret", "PF", "Avg Win", "Avg Loss"]):
    t.rows[0].cells[i].text = h
for i, v in enumerate(["4.38", "6.9%", "56.2%", "698", "83.5%", "4.46", "9.07%", "-2.46%"]):
    t.rows[1].cells[i].text = v

doc.add_paragraph(
    "Year-by-year returns: 2018: 57.8%, 2019: 29.2%, 2020: 81.9%, "
    "2021: 109.3%, 2022: 110.8%, 2023: 131.0%, 2024: 89.7%, 2025: 81.2%. "
    "Positive every year. Worst month: -2.4%. Only 6% of months negative. "
    "Max drawdown occurred 2018-12-24 (Christmas Eve selloff)."
)
doc.add_paragraph(
    "Trade profile: 87 trades/year average. 86% of exits via 3% trailing stop "
    "(600 of 698 trades). Only 95 trades held to expiry. Avg holding: 7.3 trading days. "
    "85% of losses are exactly -3.2% (the 3% trail plus 0.2% slippage). "
    "Winners average +9.07% with a fat right tail (best trade: +107%). "
    "2019 had only 11 trades due to compressed ML_90 predictions -- still profitable."
)
doc.add_paragraph(
    "Why it works: EP converts the ML directional edge into a bounded-loss strategy. "
    "When the model picks a winner (56% of the time), the trailing stop lets it run. "
    "When it picks a loser, the 3% trail limits damage. Combined with T90 selectivity "
    "and C1 diversification, the portfolio is never exposed to concentrated drawdowns."
)
doc.add_paragraph(
    "Top sectors traded: Information Technology (150), Consumer Discretionary (90), "
    "Health Care (75), Industrials (62), Financials (47). "
    "Trades per year: 2018: 80, 2019: 11, 2020: 89, 2021: 91, "
    "2022: 139, 2023: 102, 2024: 90, 2025: 96."
)

# STRATEGY 6
add_heading("Strategy #6 -- High Win Rate Cash Machine", level=3)
doc.add_paragraph(
    "Configuration: WP ranking, T90 threshold, EM exit (MFE trailing stop), "
    "Quarter Kelly sizing, 3 max positions, C1 (1 per sector)."
)
t = doc.add_table(rows=2, cols=8)
# t.style applied after creation
for i, h in enumerate(["Sharpe", "Max DD", "WR", "Trades", "Ann Ret", "PF", "Avg Win", "Avg Loss"]):
    t.rows[0].cells[i].text = h
for i, v in enumerate(["3.44", "9.9%", "85.3%", "442", "62.8%", "6.00", "7.21%", "-7.37%"]):
    t.rows[1].cells[i].text = v

doc.add_paragraph(
    "Year-by-year returns: 2018: 43.3%, 2019: 23.5%, 2020: 49.0%, "
    "2021: 89.9%, 2022: 65.1%, 2023: 96.6%, 2024: 82.0%, 2025: 67.2%. "
    "Positive every year. Narrower range than #23. "
    "Worst month: -5.6%. 15% of months negative."
)
doc.add_paragraph(
    "Trade profile: 55 trades/year. 55% exit via MFE trailing stop (244), "
    "44% held to expiry (195). Avg holding: 11.7 trading days. "
    "Losses are larger (-7.37% avg) but much less frequent (14.7%). "
    "Max drawdown occurred 2020-05-01 during COVID aftermath. "
    "2019 had only 9 trades but all captured enough for 23.5% return."
)
doc.add_paragraph(
    "Why it works: EM locks in profits once predicted return is hit, then trails. "
    "85% win rate means 5 of 6 trades profitable. Psychologically comfortable. "
    "Losses are larger per trade than EP but occur far less often."
)
doc.add_paragraph(
    "Top sectors: Information Technology (82), Health Care (52), "
    "Consumer Discretionary (51), Industrials (42), Financials (36). "
    "Trades per year: 2018: 53, 2019: 9, 2020: 63, 2021: 60, "
    "2022: 68, 2023: 63, 2024: 61, 2025: 65."
)

# STRATEGY 76
add_heading("Strategy #76 -- Ultra Conservative (2 Positions)", level=3)
doc.add_paragraph(
    "Configuration: CW composite ranking, T85 threshold, EM exit, "
    "Quarter Kelly sizing, 2 max positions, C1 (1 per sector)."
)
t = doc.add_table(rows=2, cols=8)
# t.style applied after creation
for i, h in enumerate(["Sharpe", "Max DD", "WR", "Trades", "Ann Ret", "PF", "Avg Win", "Avg Loss"]):
    t.rows[0].cells[i].text = h
for i, v in enumerate(["3.32", "10.8%", "89.3%", "307", "75.9%", "22.52", "10.77%", "-10.08%"]):
    t.rows[1].cells[i].text = v

doc.add_paragraph(
    "Year-by-year returns: 2018: 21.6%, 2019: 15.4%, 2020: 66.1%, "
    "2021: 71.1%, 2022: 77.9%, 2023: 138.8%, 2024: 169.0%, 2025: 99.9%. "
    "Strong upward trend across years. Positive every year. "
    "Longest losing streak: 2 trades. Max drawdown 2018-10-29."
)
doc.add_paragraph(
    "Trade profile: 38 trades/year. Highest win rate (89.3%) and highest profit factor "
    "(22.52) of all strategies. Only 33 losing trades in 8 years. "
    "CW composite selects the best balance of WP, predicted return, and MFE gap. "
    "62% exit via trailing stop, 37% held to expiry."
)
doc.add_paragraph(
    "Top sectors: Consumer Discretionary (61), Information Technology (50), "
    "Health Care (32), Energy (32), Financials (27). "
    "Trades per year: 2018: 38, 2019: 14, 2020: 39, 2021: 45, "
    "2022: 46, 2023: 44, 2024: 40, 2025: 41."
)

# STRATEGY 24
add_heading("Strategy #24 -- Highest Weighted Score", level=3)
doc.add_paragraph(
    "Configuration: CW composite ranking, T85 threshold, EP exit, "
    "Adaptive Kelly sizing, 4 max positions, C2 (2 per sector)."
)
t = doc.add_table(rows=2, cols=8)
# t.style applied after creation
for i, h in enumerate(["Sharpe", "Max DD", "WR", "Trades", "Ann Ret", "PF", "Avg Win", "Avg Loss"]):
    t.rows[0].cells[i].text = h
for i, v in enumerate(["4.40", "12.2%", "58.6%", "1180", "489%", "7.68", "11.62%", "-2.45%"]):
    t.rows[1].cells[i].text = v

doc.add_paragraph(
    "Year-by-year returns: 2018: 267%, 2019: 128%, 2020: 747%, "
    "2021: 210%, 2022: 526%, 2023: 848%, 2024: 1145%, 2025: 746%. "
    "Highest weighted score (0.969) due to combination of high Sharpe, "
    "moderate drawdown, and 8 years profitable."
)
doc.add_paragraph(
    "148 trades/year with Adaptive Kelly sizing means larger bets on "
    "higher-confidence signals. The CW + EP combination appears across "
    "3 of the top 5 strategies, confirming it as a robust parameter region. "
    "94% of exits via 3% trailing stop. Avg holding: 6.0 trading days."
)
doc.add_paragraph(
    "Caution: The 489% annualized return reflects aggressive compounding. "
    "In live trading, position sizes would eventually exceed practical liquidity limits. "
    "Included as evidence of the parameter region strength, not a literal return target."
)

# --- 11.5 ---
add_heading("11.5 Robust Parameter Region", level=2)
doc.add_paragraph(
    "The backtest reveals a clear robust region where multiple nearby parameter "
    "combinations all produce strong results. Changing one parameter does not "
    "destroy performance."
)
doc.add_paragraph(
    "The robust region: Quarter Kelly (SK) or Adaptive Kelly (SA) + "
    "C1 sector limits + T85 or T90 threshold + EP or EM exit."
)

t = doc.add_table(rows=8, cols=7)
# t.style applied after creation
for i, h in enumerate(["#", "Config", "Sharpe", "Max DD", "WR", "Ann Ret", "Profile"]):
    t.rows[0].cells[i].text = h
rows = [
    ["23", "WP/T90/EP/SK/P3/C1", "4.38", "6.9%", "56%", "84%", "Safest"],
    ["6", "WP/T90/EM/SK/P3/C1", "3.44", "9.9%", "85%", "63%", "High WR"],
    ["76", "CW/T85/EM/SK/P2/C1", "3.32", "10.8%", "89%", "76%", "Ultra conservative"],
    ["21", "WP/T85/EP/SK/P3/C2", "3.56", "8.3%", "59%", "98%", "EP with C2"],
    ["25", "WP/T85/EM/SA/P3/C1", "3.38", "17.6%", "85%", "120%", "Adaptive sizing"],
    ["66", "WP/T85/EM/SK/P2/C1", "2.84", "9.9%", "83%", "39%", "Min positions"],
    ["8", "WP/T90/ET/SK/P3/C1", "2.71", "12.1%", "71%", "56%", "Time+trail"],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "All 7 strategies in this region are profitable every year with max drawdown "
    "under 18%. Changing exit rule, threshold, or position count shifts the risk/return "
    "profile but does not break the strategy. This is the hallmark of a genuine edge."
)

# --- 11.6 ---
add_heading("11.6 Key Findings and Implications", level=2)

add_heading("Finding 1: EP Exit Is the Biggest Alpha Source", level=3)
doc.add_paragraph(
    "The 3% percentage trailing stop (EP) is the single most impactful parameter. "
    "It raises average Sharpe from 2.59 (all other exits) to 3.95 while cutting drawdown "
    "from 26.6% to 10.2%. EP transforms every trade into a bounded bet: max loss ~3.2% "
    "(3% trail + 0.2% slippage), uncapped upside. With 58% winning at avg 9% and 42% "
    "losing at 3.2%, expected value per trade = 0.58 * 9% - 0.42 * 3.2% = +3.9%."
)

add_heading("Finding 2: Sector Diversification Matters More Than Trade Selection", level=3)
doc.add_paragraph(
    "C1 vs CN produces a larger Sharpe improvement (+0.71) than the best vs worst "
    "ranking method (+0.53). Drawdown reduction is 16pp (17.6% vs 33.8%). "
    "Correlated sector exposure is the primary portfolio risk, not individual trade risk."
)

add_heading("Finding 3: Quarter Kelly Is the Sweet Spot", level=3)
doc.add_paragraph(
    "SK achieves 94% higher Sharpe than SF with 51% less drawdown. "
    "SH produces higher total returns but 66% more drawdown. "
    "For a cash machine prioritizing survival, SK is the clear choice."
)

add_heading("Finding 4: 2019 Is the Structural Weakness", level=3)
doc.add_paragraph(
    "All 11 strategies with a losing year lost money only in 2019. "
    "Root cause: the walk-forward model trained on 2000-2018 produces compressed predictions, "
    "yielding only 722 ML_85+ long opportunities (vs 240K-320K normally). "
    "At ML_70 there were 300K+ candidates and all T70 strategies stayed profitable. "
    "Implication: T85/T90 strategies should have a fallback -- if daily candidate count "
    "drops below ~50, temporarily lower to T70 rather than sitting out."
)

add_heading("Finding 5: Composite Rankings Beat Single Metrics", level=3)
doc.add_paragraph(
    "CR and CW outperform all single-metric rankings. CR achieves the highest "
    "average Sharpe at 3.20. No single model output is sufficient -- the best picks "
    "balance win probability with expected return magnitude and MFE runway."
)

add_heading("Finding 6: More Positions Increase Return but Also Drawdown", level=3)
doc.add_paragraph(
    "2 positions: 9.9% DD, 39-90% ann return. "
    "4 positions: 17-23% DD, 85-184% ann return. "
    "8 positions: 15-26% DD, 113-192% ann return. "
    "Marginal benefit diminishes after 4-5 positions. "
    "For a cash machine, 2-3 positions with C1 is optimal."
)

# --- 11.7 ---
add_heading("11.7 Strategy Correlations", level=2)
doc.add_paragraph(
    "Daily return correlations between top strategies range 0.51-0.72. "
    "#23 and #6 are most correlated (0.72, same base config, differ only in exit). "
    "#76 and #23 are least correlated (0.51), suggesting they could combine "
    "for further diversification. Moderate correlations indicate strategies share "
    "the underlying ML edge but express it differently enough that a multi-strategy "
    "approach would reduce portfolio-level variance."
)

# --- 11.8 ---
add_heading("11.8 Monthly Return Statistics", level=2)
t = doc.add_table(rows=6, cols=6)
# t.style applied after creation
for i, h in enumerate(["Strategy", "Median Mo", "Worst Mo", "Best Mo", "% Neg", "Mo Std"]):
    t.rows[0].cells[i].text = h
rows = [
    ["#23 (EP/T90/SK/C1)", "5.1%", "-2.4%", "20.6%", "6%", "4.1%"],
    ["#21 (EP/T85/SK/C2)", "5.4%", "-2.8%", "37.5%", "8%", "5.8%"],
    ["#6 (EM/T90/SK/C1)", "4.1%", "-5.6%", "15.0%", "15%", "4.2%"],
    ["#76 (EM/T85/SK/C1)", "4.6%", "-8.3%", "23.1%", "14%", "5.2%"],
    ["#24 (EP/T85/SA/C2)", "14.3%", "-2.6%", "58.4%", "9%", "13.1%"],
]
for r, rd in enumerate(rows):
    for c, v in enumerate(rd):
        t.rows[r + 1].cells[c].text = v

doc.add_paragraph(
    "#23 has the tightest monthly distribution: only 6% negative months, worst -2.4%. "
    "This is the most cash-machine-like profile. #6 has more negative months (15%) "
    "but bounded by high WR limiting consecutive losses."
)

# --- 11.9 ---
add_heading("11.9 Compounding and Practical Limits", level=2)
doc.add_paragraph(
    "All returns assume full compounding: each new trade allocation computed from "
    "current equity. This produces extreme total returns over 8 years. In practice, "
    "three factors limit compounding:"
)
doc.add_paragraph(
    "1. Liquidity: As equity grows, position sizes exceed daily volume of smaller "
    "S&P 500 stocks. A $50M portfolio with 3 positions at 30% needs $15M per position."
)
doc.add_paragraph(
    "2. Market impact: Large orders increase slippage beyond modeled 0.1%."
)
doc.add_paragraph(
    "3. Capacity: The ML edge may degrade if too much capital chases the same signals."
)
doc.add_paragraph(
    "For strategy comparison, relative ranking is valid regardless of compounding effects. "
    "For live capital planning, use per-trade statistics (win rate, avg win/loss, "
    "profit factor) rather than compounded total returns."
)

# --- 11.10 ---
add_heading("11.10 Recommendation for Live Trading", level=2)
doc.add_paragraph(
    "Primary strategy: #23 (WP / T90 / EP / SK / 3 positions / C1). "
    "Safest configuration: 6.9% max drawdown, 4.38 Sharpe, worst month -2.4%, "
    "positive all 8 years including structurally weak 2019."
)
doc.add_paragraph(
    "Secondary strategy: #6 (WP / T90 / EM / SK / 3 positions / C1). "
    "For traders preferring high win rate (85%) over lowest drawdown. "
    "5 of every 6 trades win."
)
doc.add_paragraph(
    "Both share the robust region (SK + C1 + T90) and differ only in exit rule. "
    "Shadow-trade both during validation phase (Section 9.2). After 50+ shadow trades, "
    "select the one whose live performance matches backtest expectations."
)
doc.add_paragraph(
    "Implementation priority: Build #23 first (EP is simpler -- just a trailing stop, "
    "no MFE prediction dependency). Add #6 as the second implementation."
)

# --- 11.11 ---
add_heading("11.11 Full Strategy Ranking", level=2)
doc.add_paragraph(
    "Complete ranking of all 112 strategies by weighted score. "
    "Weighted score = 0.40 * (years_profitable/8) + 0.25 * normalized_sharpe "
    "+ 0.20 * (1 - normalized_drawdown) + 0.15 * normalized_return."
)

# Load and add the full ranking table
import pandas as pd
summary = pd.read_csv("results/backtest/summary.csv")
summary = summary.sort_values("weighted_score", ascending=False)

# Add a compact table with all 112 strategies
cols_for_table = [
    "strategy_id", "category", "ranking", "threshold", "exit", "sizing",
    "max_positions", "concentration", "sharpe_ratio", "max_drawdown",
    "win_rate", "total_trades", "annualized_return", "years_profitable",
    "weighted_score",
]
headers = ["#", "Category", "Rank", "Thr", "Exit", "Size", "Pos", "Conc",
           "Sharpe", "MaxDD", "WR", "Trades", "AnnRet", "Yr+", "WtSc"]

t = doc.add_table(rows=len(summary) + 1, cols=len(headers))
# t.style applied after creation

# Make font smaller for this large table
for i, h in enumerate(headers):
    cell = t.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(7)

for row_idx, (_, row) in enumerate(summary.iterrows()):
    vals = [
        str(int(row["strategy_id"])),
        row["category"][:12],
        row["ranking"],
        str(int(row["threshold"])),
        row["exit"],
        row["sizing"],
        str(int(row["max_positions"])),
        row["concentration"],
        f'{row["sharpe_ratio"]:.2f}',
        f'{row["max_drawdown"]:.1%}',
        f'{row["win_rate"]:.0%}',
        str(int(row["total_trades"])),
        f'{row["annualized_return"]:.0%}',
        str(int(row["years_profitable"])),
        f'{row["weighted_score"]:.3f}',
    ]
    for col_idx, v in enumerate(vals):
        cell = t.rows[row_idx + 1].cells[col_idx]
        cell.text = v
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(7)

# --- 11.12 ---
add_heading("11.12 Year-by-Year Returns (All Strategies)", level=2)
doc.add_paragraph(
    "Year-by-year returns for all 112 strategies, sorted by weighted score. "
    "All values are annual returns (not cumulative)."
)

yr_headers = ["#", "Config"] + [str(y) for y in range(2018, 2026)] + ["Yr+"]
yr_cols = [f"year_{y}" for y in range(2018, 2026)]

t = doc.add_table(rows=len(summary) + 1, cols=len(yr_headers))
# t.style applied after creation

for i, h in enumerate(yr_headers):
    cell = t.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(7)

for row_idx, (_, row) in enumerate(summary.iterrows()):
    config = f'{row["ranking"]}/T{int(row["threshold"])}/{row["exit"]}/{row["sizing"]}/P{int(row["max_positions"])}/{row["concentration"]}'
    vals = [str(int(row["strategy_id"])), config]
    for yc in yr_cols:
        vals.append(f'{row[yc]:.0%}')
    vals.append(str(int(row["years_profitable"])))
    for col_idx, v in enumerate(vals):
        cell = t.rows[row_idx + 1].cells[col_idx]
        cell.text = v
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(6)

# --- 11.13 ---
add_heading("11.13 Output Data Files", level=2)
doc.add_paragraph(
    "All backtest output saved in results/backtest/:"
)
doc.add_paragraph(
    "summary.csv: One row per strategy (112 rows). All metrics plus weighted scores."
)
doc.add_paragraph(
    "trades.csv: One row per trade (76,709 rows). Includes strategy_id, symbol, sector, "
    "entry/exit dates, ML predictions, allocation, return, P&L, exit reason, HWM."
)
doc.add_paragraph(
    "equity.csv: One row per strategy per day (225,232 rows). "
    "Full equity curve with cash, invested, position count, drawdown."
)

# SAVE
doc.save("docs/TradeWave_Stock_Strategy_Playbook_V2.docx")
print("Document saved successfully with Section 11 (Backtest Results).")
