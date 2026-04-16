"""Add minimum account size section to Options Playbook."""
from docx import Document
from docx.oxml.ns import qn

doc = Document("docs/TradeWave_Options_Strategy_Playbook_V2.docx")

seen = set()
for child in list(doc.styles.element):
    sid = child.get(qn("w:styleId"))
    if sid and sid in seen:
        doc.styles.element.remove(child)
    elif sid:
        seen.add(sid)

def heading(text, level=1):
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    for x in pPr.findall(qn("w:pStyle")):
        pPr.remove(x)
    pPr.insert(0, pPr.makeelement(qn("w:pStyle"), {qn("w:val"): f"Heading{level}"}))

def table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r, rd in enumerate(rows):
        for c, v in enumerate(rd):
            t.rows[r+1].cells[c].text = str(v)


doc.add_page_break()
heading("13. Minimum Account Size for Automation")

doc.add_paragraph(
    "This section defines the minimum capital required to automate the options spread "
    "strategies. The key constraint is that options contracts are sold in units of 100 shares. "
    "You cannot buy half a contract. This creates a fixed minimum per trade regardless of "
    "account size."
)

heading("13.1 Per-Trade Capital Requirements", level=2)

doc.add_paragraph(
    "Spread collateral/premium scales with stock price. S&P 500 stocks range from ~$30 to $500+. "
    "Average S&P 500 stock price is approximately $150."
)

table(
    ["Structure", "$50 Stock", "$100 Stock", "$150 Stock (avg)", "$200 Stock", "$500 Stock"],
    [
        ["Credit spread (3% wide)", "$150", "$300", "$450", "$600", "$1,500"],
        ["Debit spread (3% wide, 40% cost)", "$60", "$120", "$180", "$240", "$600"],
    ],
)

doc.add_paragraph(
    "Credit spread collateral = spread_width x 100 shares per contract. "
    "Debit spread premium = cost_ratio x spread_width x 100 shares. "
    "One contract is the minimum trade size. On the average S&P 500 stock, "
    "a credit spread requires ~$450 collateral and a debit spread requires ~$180 premium."
)

heading("13.2 Account Minimums by Strategy", level=2)

table(
    ["Strategy", "Per Trade", "3 Positions", "50% Reserve", "Buffer", "Minimum", "Practical"],
    [
        ["Credit spreads only", "$450", "$1,350", "$2,700", "+$1,350", "$2,500", "$5,000"],
        ["Debit spreads only", "$180", "$540", "$1,080", "+$540", "$1,500", "$3,000"],
        ["Both strategies", "$630", "$1,890", "$3,780", "+$1,890", "$4,000", "$5,000-$10,000"],
    ],
)

doc.add_paragraph(
    "The 50% cash reserve rule from the backtester spec (Section 7) requires that total "
    "deployed premium/collateral never exceeds 50% of account equity. The buffer covers "
    "the worst case of 3 simultaneous max losses. Most brokers require $2,000 minimum "
    "for spread trading approval."
)

heading("13.3 Projected Returns by Account Size", level=2)

heading("Credit Spreads Only (55 trades/yr, 33% EV on collateral)", level=3)
table(
    ["Account", "Per-Trade Deploy", "Annual $ (backtest)", "Return %", "Conservative (-30%)", "Pessimistic (-50%)"],
    [
        ["$2,500", "$250", "$4,538", "182%", "$3,176", "$2,269"],
        ["$5,000", "$450", "$8,168", "163%", "$5,717", "$4,084"],
        ["$10,000", "$450*", "$8,168", "82%", "$5,717", "$4,084"],
        ["$25,000", "$450*", "$8,168", "33%", "$5,717", "$4,084"],
    ],
)
doc.add_paragraph(
    "* Returns flatten above ~$5K at 1 contract per trade because contract size is fixed. "
    "To scale beyond $5K, trade 2+ contracts per position (requires ~$10K for credit, "
    "~$5K for debit to maintain the same risk profile with 2 contracts)."
)

heading("Debit Spreads Only (45 trades/yr, 70% EV on premium)", level=3)
table(
    ["Account", "Per-Trade Deploy", "Annual $ (backtest)", "Return %", "Conservative (-30%)", "Pessimistic (-50%)"],
    [
        ["$1,500", "$150", "$4,725", "315%", "$3,308", "$2,363"],
        ["$3,000", "$180", "$5,670", "189%", "$3,969", "$2,835"],
        ["$5,000", "$180*", "$5,670", "113%", "$3,969", "$2,835"],
    ],
)

heading("Both Strategies Combined", level=3)
table(
    ["Account", "Annual $ (backtest)", "Return %", "Conservative (-30%)", "Pessimistic (-50%)"],
    [
        ["$5,000", "$10,208", "204%", "$7,146", "$5,104"],
        ["$10,000", "$13,838", "138%", "$9,687", "$6,919"],
        ["$25,000 (2 contracts)", "$27,676", "111%", "$19,373", "$13,838"],
    ],
)

heading("13.4 Commission Impact", level=2)

table(
    ["Account Size", "Annual Commissions*", "Drag %", "Assessment"],
    [
        ["$2,500", "~$130", "5.2%", "Significant. Use low-cost broker."],
        ["$5,000", "~$130", "2.6%", "Manageable."],
        ["$10,000", "~$130", "1.3%", "Negligible."],
        ["$25,000", "~$130", "0.5%", "Irrelevant."],
    ],
)
doc.add_paragraph(
    "* Assumes ~100 trades/yr x 2 legs x $0.65/contract (IBKR rates). "
    "TastyTrade charges $1/contract to open, $0 to close. "
    "Robinhood charges $0 but has wider spreads. "
    "For accounts under $5K, use TastyTrade or IBKR to minimize drag."
)

heading("13.5 Broker Requirements", level=2)

table(
    ["Broker", "Minimum for Spreads", "Commission/Contract", "Notes"],
    [
        ["IBKR", "$2,000", "$0.65", "Best execution, API for automation"],
        ["TastyTrade", "No minimum", "$1.00 open / $0 close", "Options-focused, good fills"],
        ["Schwab/TD", "$2,000", "$0.65", "Reliable, good platform"],
        ["Robinhood", "$2,000", "$0", "Free but wider spreads, limited API"],
    ],
)

doc.add_paragraph(
    "For automation: IBKR is recommended. It has a robust API (TWS API or IBKR Client Portal), "
    "low commissions, and good execution quality. The $2,000 minimum is easily met. "
    "TastyTrade is the alternative for manual or semi-automated trading."
)

heading("13.6 Scaling Path", level=2)

doc.add_paragraph(
    "Phase 1: Start with $5,000. Run credit spreads only (93% WR, psychologically comfortable). "
    "1 contract per trade. Target: validate live performance against backtest for 50 trades (~4 months)."
)
doc.add_paragraph(
    "Phase 2: Add debit spreads at $5K. Both strategies running. "
    "Still 1 contract per trade. Target: 50 more trades across both strategies."
)
doc.add_paragraph(
    "Phase 3: If live performance within 30% of backtest after 100 trades, scale to $10-15K "
    "and begin trading 2 contracts per position on stocks under $150. "
    "This doubles the dollar returns without changing the risk profile."
)
doc.add_paragraph(
    "Phase 4: At $25K+, trade 2-3 contracts per position across the full S&P 500 price range. "
    "At this point, the options account generates $20-28K/yr at backtest rates, "
    "$14-20K/yr conservative. Commission drag is under 1%."
)

heading("13.7 Recommendation", level=2)
doc.add_paragraph(
    "Start with $5,000 on IBKR. Run credit spreads first (highest WR, income-generating, "
    "93% of trades win). Add debit spreads after 50 trades. "
    "Conservative expectation for the first year on $5K: $7,000-$10,000 in returns (140-200%). "
    "Pessimistic expectation: $5,000-$7,500 (100-150%). "
    "Worst realistic outcome (50% degradation from backtest): $3,500-$5,000 (70-100%)."
)

doc.save("docs/TradeWave_Options_Strategy_Playbook_V2.docx")
print("Options playbook: Section 13 (Minimum Account Size) added.")
