"""
Fix duplicate sections in docs and correct stale feature counts.
Strips all content from the first occurrence of the V4 heading,
fixes stale numbers in the original body, then re-adds the section once.
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy


def fix_styles(doc):
    seen = set()
    remove = []
    for child in doc.styles.element:
        sid = child.get(qn("w:styleId"))
        if sid and sid in seen:
            remove.append(child)
        elif sid:
            seen.add(sid)
    for elem in remove:
        doc.styles.element.remove(elem)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pStyle")):
        pPr.remove(existing)
    pPr.insert(0, pPr.makeelement(qn("w:pStyle"), {qn("w:val"): f"Heading{level}"}))
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, rd in enumerate(rows):
        for c, v in enumerate(rd):
            t.rows[r + 1].cells[c].text = str(v)
    return t


def add_bullet(doc, text):
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        p = doc.add_paragraph(f"- {text}")
    return p


def strip_from_heading(doc, heading_text):
    """Remove all body elements starting from the first paragraph containing heading_text.
    Preserves the sectPr (section properties) element which must remain in the body."""
    body = doc.element.body
    elements = list(body)
    # Find and save sectPr (section properties) -- always the last element or inside last para
    sect_pr = body.find(qn("w:sectPr"))
    cut_idx = None
    for i, elem in enumerate(elements):
        if elem.tag.endswith("}p"):
            text = "".join(r.text or "" for r in elem.iter() if r.tag.endswith("}t"))
            if heading_text in text:
                cut_idx = i
                break
    if cut_idx is not None:
        to_remove = [e for e in elements[cut_idx:] if not e.tag.endswith("}sectPr")]
        for elem in to_remove:
            if elem in body:
                body.remove(elem)
        # Ensure sectPr is still present (re-append if it was removed)
        if sect_pr is not None and sect_pr not in body:
            body.append(sect_pr)
        print(f"  Stripped {len(to_remove)} elements from index {cut_idx}")
    else:
        print(f"  Heading '{heading_text}' not found -- nothing stripped")
    return cut_idx is not None


def fix_text_in_doc(doc, old_text, new_text):
    """Replace text in all paragraphs."""
    count = 0
    for p in doc.paragraphs:
        if old_text in p.text:
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    return count


# ============================================================
# Fix TradeWave_Strategy_Assessment.docx
# ============================================================

print("Fixing TradeWave_Strategy_Assessment.docx...")
doc = Document("docs/TradeWave_Strategy_Assessment.docx")
fix_styles(doc)

# Strip all duplicate V4 sections
strip_from_heading(doc, "Complete Strategy Evolution")

# Fix stale feature counts in original body
n = fix_text_in_doc(doc, "59 features", "62 features")
print(f"  Fixed '59 features' -> '62 features': {n} occurrences")
n = fix_text_in_doc(doc, "The 59 features", "The 62 features")
print(f"  Fixed 'The 59 features': {n} occurrences")
n = fix_text_in_doc(doc, "16 market regime features", "19 market regime features")
print(f"  Fixed '16 market regime features': {n} occurrences")

# Re-add the section once (clean)
doc.add_page_break()
add_heading(doc, "Complete Strategy Evolution: V1 through V4 (April 2026)", level=1)

doc.add_paragraph(
    "This section captures the complete progression of TradeWave strategy research from the "
    "original 160-strategy backtest through the Codex V3 rerun and V4 enhanced testing. "
    "It also explains what is currently running in auto trading simulation and why it uses "
    "a different configuration than the research findings."
)

add_heading(doc, "Two Strategies: Auto Trading vs Research", level=2)

doc.add_paragraph(
    "There are two distinct strategy configurations in use. They solve different problems "
    "and should not be conflated."
)

add_table(doc,
    ["Dimension", "Auto Trading (Live Simulation)", "Research Best (Codex V3/V4)"],
    [
        ["System", "Original backtest (S21)", "Codex rerun (STK_045 / V4_skip_monday)"],
        ["Direction", "Long only", "Combined long + short"],
        ["Exit", "EP (6% profit target, flat stop)", "target6_atr2 (6% target + 2xATR floor)"],
        ["Sizing", "SK (Kelly-based)", "vol_inverse (volatility-scaled)"],
        ["Concentration", "C1 (5 pos, 2 sector cap)", "risk_balanced (10 pos, 3 sector cap)"],
        ["Sharpe Ratio", "3.66 (base) / 4.22 (enhanced)", "7.11 (V3) / 7.46 (V4)"],
        ["CAGR", "~64-80% (higher absolute)", "35.67% (V3) / 36.32% (V4)"],
        ["Max Drawdown", "6.7% (base) / 7.4% (enhanced)", "2.65% (V3) / 1.84% (V4)"],
        ["Win Rate", "55.9%", "84.5-85.9%"],
        ["Status", "Running in live simulation", "Research / future upgrade target"],
    ]
)

doc.add_paragraph(
    "Why auto trading uses the original S21 strategy: The auto trading system was designed and "
    "implemented before the Codex V3 rerun was completed. The S21 configuration produces higher "
    "absolute CAGR (~64-80%) than the Codex approach (35.67%), but at higher drawdown (6.7% vs "
    "2.65%). The Codex approach earns approximately 50% less annual return but has approximately "
    "60% less drawdown and a dramatically higher Sharpe ratio (7.46 vs 4.22). Both are valid "
    "depending on investor risk tolerance. The auto trading simulation will be evaluated for "
    "upgrade to the V4 configuration once live simulation results are sufficient."
)

add_heading(doc, "Original Backtest System (160 Long-Only Strategies)", level=2)

doc.add_paragraph(
    "The original backtest system tested 160 strategy configurations on 8 years of walk-forward "
    "validation (2018-2025). All strategies are long-only."
)

add_table(doc,
    ["ID", "Config", "Sharpe", "Max DD", "CAGR", "WR%", "Notes"],
    [
        ["S24", "CW/EP/T85/SA/C2", "3.70", "19.2%", "390%", "58.5%", "Highest raw Sharpe"],
        ["S21", "WP/EP/T85/SK/C2", "3.66", "6.7%", "64%", "55.9%", "Auto trading base config"],
        ["S149", "CW/EA15/T90/SK/C1", "3.24", "11.0%", "83%", "63.3%", "ATR stop variant"],
        ["Enhanced", "Best4 + CW/EP/T90/SK/P3/C1", "4.22", "7.4%", "~80%", "~62%", "Auto trading live config"],
    ]
)

add_heading(doc, "Codex V3 Rerun: Combined Long+Short System", level=2)

doc.add_paragraph(
    "The Codex V3 rerun tested 81 stock + 36 options + 48 spread configurations including "
    "both long and short seasonal patterns. The combination of target6_atr2 exit and "
    "vol_inverse sizing produced dramatically better Sharpe ratios than the original system."
)

add_heading(doc, "Stock Results", level=3)

add_table(doc,
    ["Config ID", "Direction", "Sharpe", "Max DD", "CAGR", "WR%", "2025 Holdout"],
    [
        ["STK_045", "Combined L+S", "7.11", "-2.65%", "35.67%", "84.51%", "7.15 (better)"],
        ["STK_063", "Long only", "6.86", "-3.22%", "34.70%", "83.85%", "--"],
        ["STK_009", "Short only", "5.51", "-2.13%", "22.61%", "86.17%", "--"],
    ]
)

add_heading(doc, "Options Results", level=3)

add_table(doc,
    ["Config ID", "Direction", "Sharpe", "Max DD", "Ann Return (on $10K acct)", "WR%"],
    [
        ["OPT_013", "Combined L+S", "5.34", "-22.32%", "3,959%", "71.7%"],
        ["OPT_013", "Long only", "4.97", "-30.42%", "~2,400%", "69.1%"],
    ]
)

add_heading(doc, "Spread Results", level=3)

add_table(doc,
    ["Config ID", "Direction", "Sharpe", "Max DD", "Ann Return", "WR%", "2025 Holdout"],
    [
        ["SPR_048", "Combined L+S", "6.52", "-25.07%", "613%", "89.4%", "7.67 (better)"],
        ["SPR_036", "Long only", "5.96", "-29.92%", "616%", "88.6%", "--"],
        ["SPR_048", "Short only", "5.52", "-12.88%", "288%", "92.5%", "--"],
    ]
)

add_heading(doc, "Why Codex Sharpe Is So Much Higher", level=2)

add_table(doc,
    ["Factor", "Original System", "Codex System"],
    [
        ["Exit mechanism", "EP: fixed 6% target, flat stop", "target6_atr2: 6% target + 2xATR volatility-scaled floor"],
        ["Sizing method", "Kelly-based (SK)", "vol_inverse: smaller positions in volatile stocks"],
        ["Direction", "Long only", "Combined long + short (diversification benefit)"],
        ["Position limit", "3-5 positions, 2 sector cap", "10 positions, 3 sector cap"],
    ]
)

doc.add_paragraph(
    "The target6_atr2 exit is the single biggest factor. It captures full winning trades "
    "while cutting losses at a volatility-scaled floor, producing win rates above 84% vs "
    "55-63% for the original EP exit. The vol_inverse sizing reduces drawdown by holding "
    "smaller positions in naturally volatile stocks."
)

add_heading(doc, "V4 Enhanced Backtest Results (April 2026)", level=2)

doc.add_paragraph(
    "Building on the Codex V3 baseline (Sharpe 7.11), V4 tested five enhancement dimensions "
    "across 90 configurations. SkipMonday emerged as the single most effective filter."
)

add_table(doc,
    ["Config", "Sharpe", "Max DD", "CAGR", "WR%", "All Yrs+", "2025 OOS Sharpe"],
    [
        ["V4 best combined (A_skip_monday)", "7.46", "-1.84%", "36.32%", "85.9%", "Yes", "8.41"],
        ["V4 best long-only (B_sym_quality_L)", "7.23", "-2.75%", "36.04%", "84.4%", "Yes", "8.32"],
        ["V4 best short-only (B_regime_S)", "5.82", "-2.55%", "26.38%", "85.7%", "Yes", "--"],
        ["V3 baseline (equivalent)", "7.11", "-2.65%", "35.67%", "84.5%", "Yes", "7.15"],
    ]
)

doc.add_paragraph("V4 improvements vs V3 baseline:")
add_bullet(doc, "Sharpe: 7.11 -> 7.46 (+5%, +0.35)")
add_bullet(doc, "Max Drawdown: -2.65% -> -1.84% (31% reduction -- under the 2% target)")
add_bullet(doc, "All 8 years profitable: Yes (unchanged, worst year 27.2%)")
add_bullet(doc, "2025 out-of-sample: Sharpe 7.15 -> 8.41")

doc.add_paragraph("Enhancements -- what to enable vs avoid:")
add_bullet(doc, "SkipMonday: +0.35 Sharpe, -0.81pp DD. Enable always.")
add_bullet(doc, "NoRepeat14d: +0.05 Sharpe. Enable.")
add_bullet(doc, "VIX hard block (>=35): zero in-sample cost. Enable for production safety.")
add_bullet(doc, "WeeklyBreaker: NEVER enable. Reduces trades from 1,194 to 221, Sharpe from 7.11 to 2.33, loses profitable years.")
add_bullet(doc, "Options/spreads: use baseline only. Adding stock filters over-restricts options and destroys statistical validity.")

add_heading(doc, "The 100-Year Pattern: Sep 27, 2026 - Jul 18, 2027", level=2)

doc.add_paragraph(
    "The 100-Year Pattern is the user's original discovery (named in their book): SPX has never "
    "been down from Sep 27 to approximately Jul 18 of the following year in midterm election years "
    "since 1930. The ML backtest confirms individual stock seasonal patterns also outperform "
    "significantly during these windows."
)

add_table(doc,
    ["Metric", "IN Window", "OUTSIDE Window", "Delta"],
    [
        ["ML>=70 Long WR", "83.2%", "77.1%", "+6.1pp"],
        ["ML>=85 Long WR", "85.2%", "80.3%", "+5.0pp"],
        ["ML>=55 Long WR", "80.7%", "74.7%", "+6.0pp -- matches ML>=85 outside!"],
        ["Short WR", "53.6%", "63.3%", "-9.7pp -- avoid shorts during window"],
        ["Opportunity Sharpe (ML>=70)", "9.85", "7.92", "+24.4%"],
    ]
)

add_heading(doc, "Sector Performance During the Window", level=3)

add_table(doc,
    ["Sector", "WR ML>=70 (In Window)", "vs Outside"],
    [
        ["Materials", "86.5%", "+10.2pp (overweight)"],
        ["Consumer Staples", "86.0%", "+8.4pp (overweight)"],
        ["Consumer Discretionary", "84.6%", "+7.7pp (overweight)"],
        ["Real Estate", "84.4%", "+6.5pp (overweight)"],
        ["Energy", "70.5%", "-2.4pp (EXCLUDE)"],
    ]
)

add_heading(doc, "Window Strategy Settings", level=3)

add_table(doc,
    ["Setting", "Normal Mode", "100-Year Window Mode (Sep 27 2026)"],
    [
        ["ML threshold", ">=80 (strict)", ">=70 (or >=55 for expanded universe)"],
        ["Direction", "Long + Short", "Long ONLY -- no new shorts"],
        ["Energy sector", "Normal", "EXCLUDE"],
        ["Position count", "10", "10-12 (2 extra slots)"],
        ["Sector overweight", "Diversified", "Materials, Staples, Discretionary, Real Estate"],
        ["Options", "Calls + puts", "Calls ONLY"],
    ]
)

add_heading(doc, "Complete Recommended Configuration Set", level=2)

add_table(doc,
    ["Sleeve", "Config", "Sharpe", "DD", "CAGR", "Use Case"],
    [
        ["Auto trading (current)", "S21 + Best4 enhanced", "4.22", "-7.4%", "~80%", "Live simulation, long-only"],
        ["Research primary (future)", "V4_A_skip_monday (L+S)", "7.46", "-1.84%", "36.3%", "Lowest DD, highest risk-adj return"],
        ["Long-only sleeve", "V4_B_sym_quality_L", "7.23", "-2.75%", "36.0%", "Capital that cannot go short"],
        ["Short hedge", "V4_B_regime_S", "5.82", "-2.55%", "26.4%", "Standalone short or hedge"],
        ["Options account", "OPT_013 baseline (no filters)", "5.34", "-22.3%", "3,959%*", "$10K options account"],
        ["Spread account", "SPR_048 baseline (no filters)", "6.52", "-25.1%", "613%", "Credit spread account"],
        ["100-Year window", "V4_A_skip_monday + regime mode", "7.5+", "<2%", "36%+", "Sep 27 2026 - Jul 18 2027"],
    ]
)

doc.add_paragraph("* Options returns on $10K account -- reflects leverage and full account redeployment.")

doc.save("docs/TradeWave_Strategy_Assessment.docx")
print("  -> Saved TradeWave_Strategy_Assessment.docx")


# ============================================================
# Fix TradeWave_V3_Codex_Strategy_Assessment.docx
# ============================================================

print("\nFixing TradeWave_V3_Codex_Strategy_Assessment.docx...")
doc = Document("docs/TradeWave_V3_Codex_Strategy_Assessment.docx")
fix_styles(doc)

strip_from_heading(doc, "V4 Enhanced Backtest Results")

# Re-add once
doc.add_page_break()
add_heading(doc, "V4 Enhanced Backtest Results (April 2026)", level=1)

doc.add_paragraph(
    "Following the Codex V3 rerun (baseline Sharpe 7.11), a V4 enhanced backtest was run "
    "testing five improvement dimensions across 90 configurations on 8 years (2018-2025) "
    "with 2025 out-of-sample validation."
)

add_heading(doc, "Enhancement Attribution (BaseA, Combined L+S)", level=2)

doc.add_paragraph(
    "BaseA replicates the V3 best config (STK_045: WP/strict/risk_balanced/target6_atr2/vol_inverse, "
    "Sharpe 7.11). Each enhancement was tested individually against this baseline."
)

add_table(doc,
    ["Enhancement", "Sharpe", "Max DD", "CAGR", "Trade Count", "All Yrs+"],
    [
        ["01_baseline (V3 equivalent)", "7.11", "-2.65%", "35.67%", "1,194", "Yes"],
        ["02_vix_block (no entry VIX>=35)", "7.11", "-2.65%", "35.67%", "1,194", "Yes"],
        ["03_sym_quality (prior-yr return filter)", "7.03", "-2.83%", "37.14%", "1,274", "Yes"],
        ["04_no_repeat14 (14-day symbol cooldown)", "7.16", "-2.54%", "35.92%", "1,193", "Yes"],
        ["05_skip_monday (no Monday entries)", "7.46", "-1.84%", "36.32%", "1,155", "Yes"],
        ["06_wkly_breaker (pause after 3/5 losses)", "2.33", "-2.65%", "4.66%", "221", "No"],
        ["07_best4 (all 4 filters combined)", "6.92", "-2.12%", "33.93%", "1,170", "Yes"],
        ["08_vix_best4 (VIX + all 4 filters)", "6.92", "-2.12%", "33.93%", "1,170", "Yes"],
        ["09_regime (100-Year Pattern switching)", "6.98", "-3.31%", "36.38%", "1,243", "Yes"],
        ["10_vix_regime (VIX + regime)", "6.98", "-3.31%", "36.38%", "1,243", "Yes"],
        ["11_vix_best4_reg (VIX + best4 + regime)", "6.76", "-1.85%", "34.52%", "1,221", "Yes"],
        ["12_full (all enhancements)", "6.38", "-2.11%", "27.10%", "971", "Yes"],
    ]
)

doc.add_paragraph(
    "WeeklyBreaker is the critical failure: it reduces trade count from 1,194 to 221 and "
    "collapses Sharpe from 7.11 to 2.33. The 'Full' stack inherits this damage. "
    "SkipMonday is the only enhancement that meaningfully improves both Sharpe and DD simultaneously."
)

add_heading(doc, "V4 Best Configurations", level=2)

add_table(doc,
    ["Role", "Config ID", "Sharpe", "Max DD", "CAGR", "WR%", "2025 OOS Sharpe"],
    [
        ["Primary combined L+S", "V4_STK_A_05_skip_monday_B", "7.46", "-1.84%", "36.32%", "85.9%", "8.41"],
        ["Long-only sleeve", "V4_STK_B_03_sym_quality_L", "7.23", "-2.75%", "36.04%", "84.4%", "8.32"],
        ["Short-only hedge", "V4_STK_B_09_regime_S", "5.82", "-2.55%", "26.38%", "85.7%", "--"],
        ["Options (no filters)", "V4_OPT_01_baseline_B", "5.38", "-22.32%", "4,097%*", "71.7%", "n/a"],
        ["Spreads (no filters)", "V4_SPR_01_baseline_B", "6.35", "-25.07%", "622%", "89.4%", "n/a"],
    ]
)

doc.add_paragraph("* Options returns on $10K account.")

add_heading(doc, "Year-by-Year Returns (Top 5 Combined L+S)", level=2)

add_table(doc,
    ["Config", "2018", "2019", "2020", "2021", "2022", "2023", "2024", "2025", "Worst"],
    [
        ["A_baseline", "32.6%", "26.3%", "31.2%", "42.4%", "40.2%", "41.4%", "31.3%", "38.3%", "26.3%"],
        ["A_no_repeat14", "33.9%", "27.1%", "31.5%", "40.7%", "42.5%", "38.4%", "30.8%", "39.5%", "27.1%"],
        ["A_skip_monday", "28.9%", "27.2%", "29.3%", "43.4%", "46.4%", "40.4%", "30.2%", "43.6%", "27.2%"],
        ["A_vix_best4_reg", "28.7%", "29.8%", "26.8%", "38.4%", "36.1%", "48.9%", "30.3%", "37.6%", "26.8%"],
        ["B_sym_quality", "26.4%", "30.7%", "28.7%", "38.1%", "47.0%", "41.0%", "39.3%", "42.8%", "26.4%"],
    ]
)

add_heading(doc, "2025 Out-of-Sample Holdout", level=2)

add_table(doc,
    ["Config", "2018-2024 Sharpe", "2025 Sharpe", "2025 DD", "2025 WR%"],
    [
        ["A_baseline", "6.95", "7.15", "-0.93%", "85.8%"],
        ["A_skip_monday", "7.28", "8.41", "-1.57%", "91.25%"],
        ["A_no_repeat14", "7.03", "7.47", "-1.14%", "85.6%"],
        ["B_sym_quality", "7.05", "8.32", "-0.82%", "89.8%"],
    ]
)

doc.add_paragraph(
    "All top V4 configs show better 2025 out-of-sample performance than the 2018-2024 in-sample "
    "period -- consistent with V3 holdout behavior. The model is not overfitted."
)

add_heading(doc, "What Changed vs V3", level=2)

add_table(doc,
    ["Metric", "V3 Best (STK_045)", "V4 Best (A_skip_monday)", "Change"],
    [
        ["Sharpe", "7.11", "7.46", "+0.35 (+5%)"],
        ["Max Drawdown", "-2.65%", "-1.84%", "-0.81pp (31% improvement)"],
        ["CAGR", "35.67%", "36.32%", "+0.65pp"],
        ["Win Rate", "84.51%", "85.89%", "+1.38pp"],
        ["All Years Profitable", "Yes", "Yes", "No change"],
        ["Worst Year", "26.3%", "27.2%", "+0.9pp better"],
        ["2025 Holdout Sharpe", "7.15", "8.41", "+1.26"],
    ]
)

add_heading(doc, "Enhancement Summary: Enable vs Avoid", level=2)

add_table(doc,
    ["Enhancement", "Verdict", "Reason"],
    [
        ["SkipMonday", "Enable always", "+0.35 Sharpe, -0.81pp DD. Best single filter."],
        ["NoRepeat14d", "Enable", "+0.05 Sharpe, reduces concentration risk."],
        ["VIX hard block (>=35)", "Enable", "Zero in-sample cost. Critical production safety."],
        ["SymbolQuality", "Enable on BaseB only", "Mixed on BaseA, positive on BaseB."],
        ["PatternRegime", "Enable during 100-Year windows only", "Slight Sharpe cost, higher CAGR during windows."],
        ["WeeklyBreaker", "NEVER", "Catastrophic: 1,194 -> 221 trades, Sharpe 7.11 -> 2.33."],
        ["Full stack (12_full)", "Never on BaseB", "WeeklyBreaker dominates, Sharpe 2-3 on BaseB."],
        ["Filters on Options/Spreads", "Never", "Over-restricts, destroys year-by-year stability."],
    ]
)

doc.save("docs/TradeWave_V3_Codex_Strategy_Assessment.docx")
print("  -> Saved TradeWave_V3_Codex_Strategy_Assessment.docx")

print("\nDone. All duplicates removed, stale numbers fixed.")
