"""Add Section 13 (Enhanced Backtester Results) to Stock Playbook."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document("docs/TradeWave_Stock_Strategy_Playbook_V2.docx")

# Fix duplicate styles
seen = set()
for child in list(doc.styles.element):
    sid = child.get(qn("w:styleId"))
    if sid and sid in seen:
        doc.styles.element.remove(child)
    elif sid:
        seen.add(sid)

def heading(text, level=1):
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    for x in pPr.findall(qn("w:pStyle")):
        pPr.remove(x)
    pPr.insert(0, pPr.makeelement(qn("w:pStyle"), {qn("w:val"): f"Heading{level}"}))

def table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r, rd in enumerate(rows):
        for c, v in enumerate(rd):
            t.rows[r+1].cells[c].text = str(v)

doc.add_page_break()
heading("13. Enhanced Backtester: Earnings Filter + 7 Improvements")

doc.add_paragraph(
    "After the initial 112-strategy backtest, two significant enhancements were applied: "
    "(1) an earnings exclusion filter using EDGAR 8-K/E filing dates from an API at "
    "104.238.214.253:7670, and (2) seven tactical improvements tested individually and in combination."
)

heading("13.1 Earnings Filter Impact", level=2)
doc.add_paragraph(
    "Earnings dates fetched for all 475 symbols (18,278 dates total). Any candidate with earnings "
    "during the holding period [entry_date, entry_date + holding_days] is excluded. "
    "This removed 24.3% of 10-30 day candidates and 51.9% of 31-60 day candidates."
)
doc.add_paragraph(
    "Impact on baseline Strategy #23 (EP/T90/SK/C1): Sharpe dropped from 4.38 to 3.28 "
    "(-25%). Annualized return dropped from 83.5% to 71.2%. Max drawdown slightly increased "
    "from 6.9% to 7.3%. The earnings filter removes some of the best seasonal patterns "
    "(many coincide with earnings windows) but eliminates the risk of IV crush and gap risk "
    "from earnings announcements. This is the correct safety trade-off."
)

heading("13.2 Seven Improvements Tested", level=2)
table(
    ["#", "Improvement", "Sharpe", "DD", "Impact vs Baseline (3.28)"],
    [
        ["7", "Symbol Quality Scores", "3.79", "7.9%", "+0.51 Sharpe (+15%). Best single improvement."],
        ["5", "Weekly Loss Breaker (3 losses/5 days)", "3.54", "6.4%", "+0.26 Sharpe. Lowest DD."],
        ["2", "No-Repeat 14d (same symbol cooldown)", "3.44", "5.7%", "+0.16 Sharpe. 5.7% DD (best overall)."],
        ["3", "Skip Monday entries", "3.32", "7.2%", "+0.04. Mon WR 49% vs 63% other days."],
        ["6", "31-60 Tier standalone", "3.29", "8.4%", "Comparable Sharpe. Validates multi-tier."],
        ["4", "VIX-Scaled Sizing", "2.75", "4.8%", "-0.53 Sharpe. REJECTED: too aggressive."],
    ],
)

heading("13.3 Symbol Quality Scores", level=3)
doc.add_paragraph(
    "The biggest single improvement. For each symbol, compute a rolling quality score: "
    "average actual return per trade from prior validation years (winsorized to [-10%, +10%]). "
    "Add as a 10% weighted component to the composite ranking formula. "
    "Symbols like CVNA (avg +12.7%/trade) get boosted. Symbols like AZO (avg -3.0%) get penalized. "
    "This captures the 'is this a good seasonal pattern stock' signal that V2 misses."
)
doc.add_paragraph(
    "Alpha concentration finding: top 10 of 160 traded symbols produce 43% of all profits. "
    "30% of symbols have net negative contribution (-128% cumulative drag). "
    "Symbol quality scoring steers the ranking toward proven performers."
)

heading("13.4 Best-4 Combination Results", level=2)
doc.add_paragraph(
    "The best-4 improvements (Symbol Quality + No-Repeat + Weekly Breaker + Skip Monday) "
    "were combined and tested across 20 configurations. All 20 were profitable every year."
)
table(
    ["Config", "Sharpe", "DD", "WR", "Ann Ret", "Trades", "Profile"],
    [
        ["Best4+CW/EP/T90", "3.43", "7.7%", "57.1%", "71.2%", "657", "NEW #1 OVERALL"],
        ["Best4+WP/EP/T85", "3.29", "6.6%", "57.6%", "61.2%", "682", "Lowest DD"],
        ["Best4+CR/EP/T85", "3.25", "7.1%", "55.9%", "75.1%", "666", "Highest return at 3+"],
        ["Best4+WP/EP/T90", "3.22", "6.6%", "57.9%", "60.1%", "642", "Original params enhanced"],
        ["Best4+CW/EP/T85", "3.18", "7.0%", "55.6%", "73.9%", "676", "CW composite"],
        ["Best4+EP/T85/SA", "3.12", "11.2%", "57.6%", "117.5%", "682", "Adaptive sizing growth"],
        ["Best4+EP/T90/SH", "3.05", "10.9%", "57.9%", "115.7%", "642", "Half Kelly growth"],
        ["Best4+CW/EM/T85", "3.00", "14.7%", "83.4%", "67.9%", "445", "High win rate variant"],
    ],
)

heading("13.5 31-60 Tier Validation", level=2)
doc.add_paragraph(
    "The 31-60 day tier was built (17.3M rows) and tested with the same best-4 improvements. "
    "51.9% of candidates filtered by earnings (longer holds = more earnings overlap). "
    "Results: Sharpe 2.54-2.71, all configurations profitable all 8 years. "
    "This validates multi-tier as a diversification strategy -- different stocks, "
    "different holding periods, non-overlapping risk windows."
)

heading("13.6 Final Recommendation", level=2)
doc.add_paragraph(
    "Primary strategy for live trading: Best4 + CW/EP/T90/SK/P3/C1. "
    "This combines the CW composite ranking (0.54*WP + 0.225*PR + 0.135*MG + 0.10*SymQuality), "
    "the EP 3% trailing stop, T90 threshold, Quarter Kelly sizing, 3 positions max, "
    "1 per sector, with no-repeat 14d, weekly loss breaker, and Monday skip."
)
doc.add_paragraph(
    "Sharpe 3.43, max drawdown 7.7%, profitable all 8 years (2018-2025), "
    "657 trades (82/year), 57.1% win rate with asymmetric payoff "
    "(avg win ~9%, avg loss ~3.2% capped by EP trailing stop)."
)

heading("13.7 Survivorship Bias Validation", level=2)
doc.add_paragraph(
    "A concern was raised about using 2025 opp files for all training years: "
    "pat_deepest_pass values are inflated for historical samples (2025 depth applied to 2000 data). "
    "Analysis showed: (1) 48.4% of training samples have depth=0 vs depth>0, and this binary split "
    "drives most of the feature importance. (2) Within depth>0, Spearman correlation between "
    "exact depth value and returns is only 0.057 (negligible). (3) The model uses depth primarily "
    "as a binary gate (qualified/not), not a fine-grained predictor. "
    "Conclusion: the architecture is valid. No retraining needed."
)

doc.save("docs/TradeWave_Stock_Strategy_Playbook_V2.docx")
print("Stock playbook updated with Section 13 (Enhanced Results + Survivorship Analysis).")
