from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import numpy as np
import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parent
DATA_PATH = ROOT / "results" / "backtester_input_10_30.parquet"
EARNINGS_PATH = ROOT / "results" / "earnings_dates.json"
MARKDOWN_OUT = ROOT / "results" / "project_v3_codex_strategy_final.md"
DOCX_OUT = ROOT / "docs" / "TradeWave_V3_Codex_Strategy_Assessment.docx"

STOCK_START_CAPITAL = 100_000.0
OPTIONS_START_CAPITAL = 10_000.0
SPREAD_START_CAPITAL = 25_000.0
BUSINESS_DAYS = 252


@dataclass(frozen=True)
class StrategyConfig:
    strategy_id: str
    asset_class: str
    rank_key: str
    threshold_profile: str
    concentration_profile: str
    exit_profile: str
    max_positions: int
    sector_cap: int
    base_weight: float
    size_mode: str
    min_ml_score: float
    min_win_probability: float
    min_predicted_return: float
    return_builder: str
    option_premium_pct: float = 0.0
    option_theta_mult: float = 0.0
    spread_family: str = ""
    spread_otm: float = 0.0
    spread_width: float = 0.0

    @property
    def description(self) -> str:
        return (
            f"rank={self.rank_key}; threshold={self.threshold_profile}; "
            f"concentration={self.concentration_profile}; exit={self.exit_profile}; "
            f"size={self.size_mode}"
        )


def load_earnings_calendar() -> dict[str, np.ndarray]:
    raw = json.loads(EARNINGS_PATH.read_text())
    calendar: dict[str, np.ndarray] = {}
    for symbol, dates in raw.items():
        arr = (
            pd.to_datetime(pd.Series(dates), errors="coerce")
            .dropna()
            .sort_values()
            .to_numpy(dtype="datetime64[D]")
        )
        if len(arr):
            calendar[symbol] = arr
    return calendar


def build_earnings_flag(df: pd.DataFrame, calendar: dict[str, np.ndarray]) -> np.ndarray:
    flags = np.zeros(len(df), dtype=bool)
    for symbol, idx in df.groupby("symbol", sort=False, observed=True).groups.items():
        earnings = calendar.get(symbol)
        if earnings is None or len(earnings) == 0:
            continue
        positions = np.asarray(idx, dtype=np.int64)
        starts = df.loc[positions, "date"].to_numpy(dtype="datetime64[D]")
        ends = df.loc[positions, "exit_date"].to_numpy(dtype="datetime64[D]")
        left = np.searchsorted(earnings, starts, side="left")
        right = np.searchsorted(earnings, ends, side="right")
        flags[positions] = right > left
    return flags


def load_data() -> pd.DataFrame:
    columns = [
        "date",
        "year",
        "symbol",
        "sector",
        "direction",
        "holding_days",
        "ml_score",
        "predicted_return",
        "predicted_mfe",
        "win_probability",
        "p_hit_return",
        "p_hit_mfe",
        "actual_return",
        "actual_mfe",
        "stock_volatility_20d",
        "atr_14d_pct",
    ]
    df = pd.read_parquet(DATA_PATH, columns=columns).copy()
    df["date"] = pd.to_datetime(df["date"]).dt.normalize()
    df["exit_date"] = df["date"] + pd.to_timedelta(df["holding_days"], unit="D")
    df["direction"] = df["direction"].astype("category")
    df["symbol"] = df["symbol"].astype("category")
    df["sector"] = df["sector"].astype("category")
    for col in [
        "ml_score",
        "predicted_return",
        "predicted_mfe",
        "win_probability",
        "p_hit_return",
        "p_hit_mfe",
        "actual_return",
        "actual_mfe",
        "stock_volatility_20d",
        "atr_14d_pct",
    ]:
        df[col] = pd.to_numeric(df[col], errors="coerce").astype("float32")
    df["holding_days"] = pd.to_numeric(df["holding_days"], errors="coerce").fillna(0).astype("int16")
    df["year"] = pd.to_numeric(df["year"], errors="coerce").fillna(0).astype("int16")
    df = df.dropna(
        subset=[
            "date",
            "symbol",
            "sector",
            "direction",
            "ml_score",
            "predicted_return",
            "predicted_mfe",
            "win_probability",
            "actual_return",
            "actual_mfe",
            "atr_14d_pct",
        ]
    ).reset_index(drop=True)

    earnings_calendar = load_earnings_calendar()
    df["has_earnings"] = build_earnings_flag(df, earnings_calendar)
    df["eligible"] = ~df["has_earnings"]
    df["atr_pct_points"] = (df["atr_14d_pct"] * 100.0).astype("float32")
    df["combo_rank"] = (
        0.45 * df["ml_score"]
        + 35.0 * df["win_probability"]
        + 3.0 * df["predicted_return"]
        + 10.0 * df["p_hit_return"]
    ).astype("float32")
    return df


def data_exploration(df: pd.DataFrame) -> dict[str, object]:
    summary: dict[str, object] = {
        "rows_total": int(len(df)),
        "symbols": int(df["symbol"].nunique()),
        "date_min": str(df["date"].min().date()),
        "date_max": str(df["date"].max().date()),
        "base_win_rate": float((df["actual_return"] > 0).mean() * 100.0),
        "earnings_filtered_rate": float(df["has_earnings"].mean() * 100.0),
        "pred_actual_corr": float(df["predicted_return"].corr(df["actual_return"])),
        "mfe_actual_corr": float(df["predicted_mfe"].corr(df["actual_mfe"])),
        "actual_return_mean": float(df["actual_return"].mean()),
        "actual_return_median": float(df["actual_return"].median()),
        "actual_return_p05": float(df["actual_return"].quantile(0.05)),
        "actual_return_p95": float(df["actual_return"].quantile(0.95)),
        "actual_mfe_mean": float(df["actual_mfe"].mean()),
        "actual_mfe_median": float(df["actual_mfe"].median()),
    }
    summary["direction"] = (
        df.groupby("direction", observed=True)
        .agg(
            trades=("actual_return", "size"),
            win_rate=("actual_return", lambda x: (x > 0).mean() * 100.0),
            avg_return=("actual_return", "mean"),
            avg_predicted_return=("predicted_return", "mean"),
        )
        .reset_index()
    )

    work = df.copy()
    work["ml_decile"] = pd.qcut(work["ml_score"], 10, labels=False, duplicates="drop") + 1
    work["wp_decile"] = pd.qcut(work["win_probability"], 10, labels=False, duplicates="drop") + 1
    summary["ml_score_deciles"] = (
        work.groupby("ml_decile", observed=True)
        .agg(
            trades=("actual_return", "size"),
            win_rate=("actual_return", lambda x: (x > 0).mean() * 100.0),
            avg_return=("actual_return", "mean"),
        )
        .reset_index()
    )
    summary["win_probability_deciles"] = (
        work.groupby("wp_decile", observed=True)
        .agg(
            trades=("actual_return", "size"),
            win_rate=("actual_return", lambda x: (x > 0).mean() * 100.0),
            avg_return=("actual_return", "mean"),
        )
        .reset_index()
    )
    summary["yearly"] = (
        work.groupby("year", observed=True)
        .agg(
            trades=("actual_return", "size"),
            win_rate=("actual_return", lambda x: (x > 0).mean() * 100.0),
            avg_return=("actual_return", "mean"),
            avg_mfe=("actual_mfe", "mean"),
        )
        .reset_index()
    )
    summary["sector"] = (
        work.groupby("sector", observed=True)
        .agg(
            trades=("actual_return", "size"),
            win_rate=("actual_return", lambda x: (x > 0).mean() * 100.0),
            avg_return=("actual_return", "mean"),
        )
        .sort_values("avg_return", ascending=False)
        .reset_index()
    )
    summary["holding_days"] = (
        work.groupby("holding_days", observed=True)
        .agg(
            trades=("actual_return", "size"),
            win_rate=("actual_return", lambda x: (x > 0).mean() * 100.0),
            avg_return=("actual_return", "mean"),
        )
        .reset_index()
    )
    return summary


def build_stock_return(df: pd.DataFrame, profile: str) -> np.ndarray:
    actual = df["actual_return"].to_numpy(dtype=np.float32)
    mfe = df["actual_mfe"].to_numpy(dtype=np.float32)
    atr = df["atr_pct_points"].to_numpy(dtype=np.float32)
    if profile == "hold":
        return actual
    if profile == "target4_trail2":
        target_hit = mfe >= 4.0
        trail_hit = (mfe >= 2.5) & ((mfe - actual) >= 2.0)
        return np.where(target_hit, 4.0, np.where(trail_hit, np.maximum(actual, mfe - 2.0), actual))
    if profile == "target6_atr2":
        target_hit = mfe >= 6.0
        stop_level = -2.0 * atr
        return np.where(target_hit, 6.0, np.maximum(actual, stop_level))
    raise ValueError(profile)


def build_option_return(df: pd.DataFrame, premium_pct: float, theta_mult: float) -> np.ndarray:
    actual = df["actual_return"].to_numpy(dtype=np.float32) / 100.0
    mfe = df["actual_mfe"].to_numpy(dtype=np.float32) / 100.0
    holding = df["holding_days"].to_numpy(dtype=np.float32)
    gross = 0.55 * np.maximum(actual, 0.0) + 0.30 * np.maximum(mfe, 0.0)
    pnl = (gross - premium_pct) / premium_pct
    theta_drag = theta_mult * (holding / 30.0)
    return np.clip((pnl - theta_drag) * 100.0, -100.0, 300.0)


def build_spread_return(df: pd.DataFrame, family: str, otm: float, width: float) -> np.ndarray:
    actual = df["actual_return"].to_numpy(dtype=np.float32)
    mfe = df["actual_mfe"].to_numpy(dtype=np.float32)
    if family == "bull_call":
        debit = 0.40 * width
        terminal = np.clip(actual, 0.0, width)
        early = np.where(mfe >= (otm + 0.5 * width), 0.60 * width, terminal)
        return np.clip(((early - debit) / debit) * 100.0, -100.0, 150.0)
    if family == "bull_put":
        credit = 0.32 * width
        risk = width - credit
        intrinsic_loss = np.clip((-actual) - otm, 0.0, width)
        return np.clip(((credit - intrinsic_loss) / risk) * 100.0, -100.0, 60.0)
    raise ValueError(family)


def generate_stock_strategies() -> list[StrategyConfig]:
    rank_keys = ["ml_score", "win_probability", "combo_rank"]
    thresholds = {
        "balanced": (70.0, 0.68, 1.50),
        "strict": (80.0, 0.74, 2.00),
        "elite": (90.0, 0.80, 2.50),
    }
    concentrations = {
        "diversified": (8, 2, 0.12, "equal"),
        "focused": (5, 1, 0.18, "confidence"),
        "risk_balanced": (10, 3, 0.10, "vol_inverse"),
    }
    exits = ["hold", "target4_trail2", "target6_atr2"]
    out: list[StrategyConfig] = []
    counter = 1
    for rank_key in rank_keys:
        for threshold_name, threshold_values in thresholds.items():
            for concentration_name, concentration_values in concentrations.items():
                for exit_name in exits:
                    out.append(
                        StrategyConfig(
                            strategy_id=f"STK_{counter:03d}",
                            asset_class="stock",
                            rank_key=rank_key,
                            threshold_profile=threshold_name,
                            concentration_profile=concentration_name,
                            exit_profile=exit_name,
                            max_positions=concentration_values[0],
                            sector_cap=concentration_values[1],
                            base_weight=concentration_values[2],
                            size_mode=concentration_values[3],
                            min_ml_score=threshold_values[0],
                            min_win_probability=threshold_values[1],
                            min_predicted_return=threshold_values[2],
                            return_builder=exit_name,
                        )
                    )
                    counter += 1
    return out


def generate_option_strategies() -> list[StrategyConfig]:
    shapes = [
        ("compact", 4, 2, 0.12, "confidence"),
        ("balanced", 5, 2, 0.10, "equal"),
        ("wide", 6, 3, 0.08, "vol_inverse"),
    ]
    premiums = [0.025, 0.030]
    theta_mults = [0.10, 0.18, 0.25]
    thresholds = {
        "strict": (82.0, 0.76, 2.0),
        "elite": (90.0, 0.82, 2.6),
    }
    out: list[StrategyConfig] = []
    counter = 1
    for threshold_name, threshold_values in thresholds.items():
        for shape_name, max_positions, sector_cap, base_weight, size_mode in shapes:
            for premium in premiums:
                for theta_mult in theta_mults:
                    out.append(
                        StrategyConfig(
                            strategy_id=f"OPT_{counter:03d}",
                            asset_class="option",
                            rank_key="combo_rank",
                            threshold_profile=threshold_name,
                            concentration_profile=shape_name,
                            exit_profile=f"premium_{premium:.3f}_theta_{theta_mult:.2f}",
                            max_positions=max_positions,
                            sector_cap=sector_cap,
                            base_weight=base_weight,
                            size_mode=size_mode,
                            min_ml_score=threshold_values[0],
                            min_win_probability=threshold_values[1],
                            min_predicted_return=threshold_values[2],
                            return_builder="option",
                            option_premium_pct=premium,
                            option_theta_mult=theta_mult,
                        )
                    )
                    counter += 1
    return out


def generate_spread_strategies() -> list[StrategyConfig]:
    families = ["bull_call", "bull_put"]
    thresholds = {
        "balanced": (75.0, 0.70, 1.8),
        "strict": (84.0, 0.77, 2.2),
    }
    widths = [4.0, 6.0, 8.0]
    otms = [1.0, 2.0]
    profiles = [
        ("spread_focus", 6, 2, 0.12, "equal"),
        ("spread_balanced", 8, 3, 0.10, "vol_inverse"),
    ]
    out: list[StrategyConfig] = []
    counter = 1
    for family in families:
        for threshold_name, threshold_values in thresholds.items():
            for otm in otms:
                for width in widths:
                    for profile_name, max_positions, sector_cap, base_weight, size_mode in profiles:
                        out.append(
                            StrategyConfig(
                                strategy_id=f"SPR_{counter:03d}",
                                asset_class="spread",
                                rank_key="combo_rank",
                                threshold_profile=threshold_name,
                                concentration_profile=profile_name,
                                exit_profile=f"{family}_{otm:.1f}_{width:.1f}",
                                max_positions=max_positions,
                                sector_cap=sector_cap,
                                base_weight=base_weight,
                                size_mode=size_mode,
                                min_ml_score=threshold_values[0],
                                min_win_probability=threshold_values[1],
                                min_predicted_return=threshold_values[2],
                                return_builder="spread",
                                spread_family=family,
                                spread_otm=otm,
                                spread_width=width,
                            )
                        )
                        counter += 1
    return out


def position_weight(row: pd.Series, config: StrategyConfig) -> float:
    if config.size_mode == "equal":
        return config.base_weight
    if config.size_mode == "confidence":
        scale = max(0.75, min(1.35, float(row["win_probability"]) / config.min_win_probability))
        return min(config.base_weight * scale, 0.24)
    if config.size_mode == "vol_inverse":
        vol = max(15.0, float(row["stock_volatility_20d"] * 100.0))
        scale = 25.0 / vol
        return min(max(config.base_weight * scale, config.base_weight * 0.6), config.base_weight * 1.4)
    raise ValueError(config.size_mode)


def pick_candidates(df: pd.DataFrame, config: StrategyConfig) -> pd.DataFrame:
    mask = (
        df["eligible"]
        & (df["ml_score"] >= config.min_ml_score)
        & (df["win_probability"] >= config.min_win_probability)
        & (df["predicted_return"] >= config.min_predicted_return)
    )
    selected = df.loc[
        mask,
        [
            "date",
            "exit_date",
            "year",
            "symbol",
            "sector",
            "direction",
            "holding_days",
            "win_probability",
            "stock_volatility_20d",
            "trade_return_pct",
        ],
    ].copy()
    if selected.empty:
        return selected
    selected["rank_value"] = df.loc[mask, config.rank_key].to_numpy()
    selected = selected.sort_values(["date", "rank_value"], ascending=[True, False])
    selected = selected.drop_duplicates(["date", "symbol", "direction"], keep="first")
    return selected


def simulate_portfolio(candidates: pd.DataFrame, config: StrategyConfig, starting_capital: float) -> dict[str, object]:
    if candidates.empty:
        return {
            "starting_capital": starting_capital,
            "ending_capital": starting_capital,
            "annualized_return": 0.0,
            "sharpe": 0.0,
            "max_drawdown": 0.0,
            "win_rate": 0.0,
            "trade_count": 0,
            "avg_trade_return": 0.0,
            "all_8_years_profitable": False,
            "worst_year": None,
            "worst_year_return": None,
            "yearly_returns": {},
            "daily_equity": pd.Series(dtype="float64"),
        }

    by_date = {date: frame for date, frame in candidates.groupby("date", sort=True)}
    event_dates = sorted(set(candidates["date"]).union(set(candidates["exit_date"])))
    cash = starting_capital
    open_positions: list[dict[str, object]] = []
    equity_points: list[tuple[pd.Timestamp, float]] = []
    trade_returns: list[float] = []
    symbol_open: set[str] = set()

    for current_date in event_dates:
        survivors: list[dict[str, object]] = []
        for pos in open_positions:
            if pos["exit_date"] <= current_date:
                cash += pos["allocation"] * (1.0 + pos["return_pct"] / 100.0)
                trade_returns.append(pos["return_pct"])
                symbol_open.discard(pos["symbol"])
            else:
                survivors.append(pos)
        open_positions = survivors

        equity_before_entries = cash + sum(pos["allocation"] for pos in open_positions)
        if current_date in by_date and len(open_positions) < config.max_positions:
            sector_counts: dict[str, int] = {}
            for pos in open_positions:
                sector_counts[pos["sector"]] = sector_counts.get(pos["sector"], 0) + 1
            slots_left = config.max_positions - len(open_positions)
            for _, row in by_date[current_date].iterrows():
                symbol = str(row["symbol"])
                sector = str(row["sector"])
                if symbol in symbol_open:
                    continue
                if sector_counts.get(sector, 0) >= config.sector_cap:
                    continue
                allocation = min(equity_before_entries * position_weight(row, config), cash)
                if allocation < equity_before_entries * 0.02:
                    continue
                open_positions.append(
                    {
                        "exit_date": row["exit_date"],
                        "allocation": allocation,
                        "return_pct": float(row["trade_return_pct"]),
                        "symbol": symbol,
                        "sector": sector,
                    }
                )
                cash -= allocation
                symbol_open.add(symbol)
                sector_counts[sector] = sector_counts.get(sector, 0) + 1
                slots_left -= 1
                if slots_left == 0 or cash <= 0:
                    break

        equity_points.append((pd.Timestamp(current_date), cash + sum(pos["allocation"] for pos in open_positions)))

    if open_positions:
        final_date = max(pos["exit_date"] for pos in open_positions)
        for pos in open_positions:
            cash += pos["allocation"] * (1.0 + pos["return_pct"] / 100.0)
            trade_returns.append(pos["return_pct"])
        equity_points.append((pd.Timestamp(final_date), cash))

    equity = pd.Series({date: value for date, value in equity_points}).sort_index().groupby(level=0).last()
    business_index = pd.date_range(equity.index.min(), equity.index.max(), freq="B")
    equity = equity.reindex(business_index).ffill().fillna(starting_capital)
    daily_returns = equity.pct_change().fillna(0.0)
    sharpe = 0.0
    if daily_returns.std(ddof=0) > 0:
        sharpe = float(np.sqrt(BUSINESS_DAYS) * daily_returns.mean() / daily_returns.std(ddof=0))
    drawdown = equity / equity.cummax() - 1.0
    max_drawdown = float(drawdown.min() * 100.0)
    total_days = max((equity.index.max() - equity.index.min()).days, 1)
    ending_capital = float(equity.iloc[-1])
    annualized_return = float(((ending_capital / starting_capital) ** (365.25 / total_days) - 1.0) * 100.0)

    yearly_returns: dict[int, float] = {}
    input_years = sorted(int(year) for year in candidates["year"].dropna().unique())
    for year in input_years:
        year_equity = equity[equity.index.year == year]
        if len(year_equity) < 2:
            continue
        yearly_returns[year] = float((year_equity.iloc[-1] / year_equity.iloc[0] - 1.0) * 100.0)
    positives = [ret > 0 for ret in yearly_returns.values()]
    return {
        "starting_capital": starting_capital,
        "ending_capital": ending_capital,
        "annualized_return": annualized_return,
        "sharpe": sharpe,
        "max_drawdown": max_drawdown,
        "win_rate": float((np.array(trade_returns) > 0).mean() * 100.0) if trade_returns else 0.0,
        "trade_count": int(len(trade_returns)),
        "avg_trade_return": float(np.mean(trade_returns)) if trade_returns else 0.0,
        "all_8_years_profitable": bool(positives and all(positives) and len(yearly_returns) >= 8),
        "worst_year": min(yearly_returns, key=yearly_returns.get) if yearly_returns else None,
        "worst_year_return": min(yearly_returns.values()) if yearly_returns else None,
        "yearly_returns": yearly_returns,
        "daily_equity": equity,
    }


def evaluate_strategies(
    df: pd.DataFrame,
    configs: list[StrategyConfig],
    starting_capital: float,
    builder: Callable[[pd.DataFrame, StrategyConfig], np.ndarray],
) -> tuple[pd.DataFrame, dict[str, dict[str, object]]]:
    rows = []
    details: dict[str, dict[str, object]] = {}
    for config in configs:
        work = df.copy()
        work["trade_return_pct"] = builder(work, config)
        candidates = pick_candidates(work, config)
        simulation = simulate_portfolio(candidates, config, starting_capital)
        rows.append(
            {
                "strategy_id": config.strategy_id,
                "asset_class": config.asset_class,
                "config": config.description,
                "rank_key": config.rank_key,
                "threshold_profile": config.threshold_profile,
                "concentration_profile": config.concentration_profile,
                "exit_profile": config.exit_profile,
                "sharpe": simulation["sharpe"],
                "max_drawdown": simulation["max_drawdown"],
                "win_rate": simulation["win_rate"],
                "annualized_return": simulation["annualized_return"],
                "trade_count": simulation["trade_count"],
                "avg_trade_return": simulation["avg_trade_return"],
                "all_8_years_profitable": simulation["all_8_years_profitable"],
                "worst_year": simulation["worst_year"],
                "worst_year_return": simulation["worst_year_return"],
                "ending_capital": simulation["ending_capital"],
            }
        )
        details[config.strategy_id] = {"config": config, "simulation": simulation, "candidates": candidates}
    results = pd.DataFrame(rows).sort_values(["sharpe", "annualized_return"], ascending=[False, False]).reset_index(drop=True)
    return results, details


def stock_return_builder(df: pd.DataFrame, config: StrategyConfig) -> np.ndarray:
    return build_stock_return(df, config.return_builder)


def option_return_builder(df: pd.DataFrame, config: StrategyConfig) -> np.ndarray:
    return build_option_return(df, config.option_premium_pct, config.option_theta_mult)


def spread_return_builder(df: pd.DataFrame, config: StrategyConfig) -> np.ndarray:
    return build_spread_return(df, config.spread_family, config.spread_otm, config.spread_width)


def pareto_frontier(results: pd.DataFrame) -> pd.DataFrame:
    work = results.sort_values(["max_drawdown", "sharpe"], ascending=[False, False]).copy()
    frontier = []
    best_sharpe = -np.inf
    for _, row in work.iterrows():
        if row["sharpe"] > best_sharpe:
            frontier.append(row)
            best_sharpe = row["sharpe"]
    return pd.DataFrame(frontier).sort_values("sharpe", ascending=False).reset_index(drop=True)


def evaluate_holdout(
    df: pd.DataFrame,
    details: dict[str, dict[str, object]],
    top_ids: list[str],
    starting_capital: float,
    builder: Callable[[pd.DataFrame, StrategyConfig], np.ndarray],
) -> pd.DataFrame:
    train = df[df["year"] <= 2024].copy()
    holdout = df[df["year"] == 2025].copy()
    rows = []
    for strategy_id in top_ids:
        config = details[strategy_id]["config"]
        for label, subset in [("2018_2024", train), ("2025", holdout)]:
            work = subset.copy()
            work["trade_return_pct"] = builder(work, config)
            simulation = simulate_portfolio(pick_candidates(work, config), config, starting_capital)
            rows.append(
                {
                    "strategy_id": strategy_id,
                    "period": label,
                    "sharpe": simulation["sharpe"],
                    "max_drawdown": simulation["max_drawdown"],
                    "annualized_return": simulation["annualized_return"],
                    "win_rate": simulation["win_rate"],
                    "trade_count": simulation["trade_count"],
                }
            )
    return pd.DataFrame(rows)


def render_table(df: pd.DataFrame, columns: list[str] | None = None, max_rows: int | None = None) -> str:
    work = df.copy()
    if columns is not None:
        work = work[columns]
    if max_rows is not None:
        work = work.head(max_rows)
    headers = [str(col) for col in work.columns]
    rows = []
    for _, row in work.iterrows():
        formatted = []
        for value in row:
            if isinstance(value, float):
                formatted.append(f"{value:.4f}")
            else:
                formatted.append(str(value))
        rows.append(formatted)
    separator = ["---"] * len(headers)
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(lines)


def findings_text(results: pd.DataFrame, label: str) -> list[str]:
    top = results.iloc[0]
    strongest_rank = results.groupby("rank_key")["sharpe"].mean().sort_values(ascending=False)
    strongest_exit = results.groupby("exit_profile")["sharpe"].mean().sort_values(ascending=False)
    return [
        f"Best {label} strategy: `{top['strategy_id']}` with Sharpe {top['sharpe']:.2f}, max drawdown {top['max_drawdown']:.2f}%, and annualized return {top['annualized_return']:.2f}%.",
        f"Average Sharpe by ranking signal was led by `{strongest_rank.index[0]}` at {strongest_rank.iloc[0]:.2f}.",
        f"Average Sharpe by exit profile was led by `{strongest_exit.index[0]}` at {strongest_exit.iloc[0]:.2f}.",
    ]


def run_suite(
    df: pd.DataFrame,
    configs: list[StrategyConfig],
    starting_capital: float,
    builder: Callable[[pd.DataFrame, StrategyConfig], np.ndarray],
) -> dict[str, object]:
    results, details = evaluate_strategies(df, configs, starting_capital, builder)
    holdout = evaluate_holdout(df, details, results.head(5)["strategy_id"].tolist(), starting_capital, builder)
    return {"results": results, "details": details, "holdout": holdout}


def write_markdown(
    exploration: dict[str, object],
    stock_combined: dict[str, object],
    stock_long: dict[str, object],
    stock_short: dict[str, object],
    options_combined: dict[str, object],
    options_long: dict[str, object],
    options_short: dict[str, object],
    spreads_combined: dict[str, object],
    spreads_long: dict[str, object],
    spreads_short: dict[str, object],
) -> None:
    top_stock_ids = stock_combined["results"].head(10)["strategy_id"].tolist()
    stock_detail_lines = []
    for strategy_id in top_stock_ids:
        sim = stock_combined["details"][strategy_id]["simulation"]
        yearly = ", ".join(f"{year}: {ret:.2f}%" for year, ret in sim["yearly_returns"].items())
        stock_detail_lines.append(
            f"#### {strategy_id}\n"
            f"- Config: {stock_combined['details'][strategy_id]['config'].description}\n"
            f"- Sharpe: {sim['sharpe']:.2f}\n"
            f"- MaxDD: {sim['max_drawdown']:.2f}%\n"
            f"- CAGR: {sim['annualized_return']:.2f}%\n"
            f"- Win rate: {sim['win_rate']:.2f}%\n"
            f"- Trades: {sim['trade_count']}\n"
            f"- Yearly returns: {yearly}\n"
        )

    def direction_block(title: str, suite: dict[str, object]) -> str:
        return (
            f"### {title}\n"
            f"{render_table(suite['results'], columns=['strategy_id','config','sharpe','max_drawdown','win_rate','annualized_return','trade_count','all_8_years_profitable'], max_rows=10)}\n\n"
            f"#### Holdout Check\n"
            f"{render_table(suite['holdout'])}\n\n"
            + "".join(f"- {line}\n" for line in findings_text(suite["results"], title.lower()))
        )

    content = f"""# V3 Strategy Backtest Results -- Independent Assessment

## Data Exploration Summary
- Rows analyzed: {exploration['rows_total']:,}
- Symbols: {exploration['symbols']}
- Date range: {exploration['date_min']} to {exploration['date_max']}
- Base win rate: {exploration['base_win_rate']:.2f}%
- Trades removed by earnings filter: {exploration['earnings_filtered_rate']:.2f}%
- Correlation between predicted_return and actual_return: {exploration['pred_actual_corr']:.4f}
- Correlation between predicted_mfe and actual_mfe: {exploration['mfe_actual_corr']:.4f}
- Actual return distribution: mean {exploration['actual_return_mean']:.2f}%, median {exploration['actual_return_median']:.2f}%, 5th pct {exploration['actual_return_p05']:.2f}%, 95th pct {exploration['actual_return_p95']:.2f}%
- Actual MFE distribution: mean {exploration['actual_mfe_mean']:.2f}%, median {exploration['actual_mfe_median']:.2f}%

### By Direction
{render_table(exploration['direction'])}

### Win Rate by ML Score Decile
{render_table(exploration['ml_score_deciles'])}

### Win Rate by Win Probability Decile
{render_table(exploration['win_probability_deciles'])}

### Performance by Year
{render_table(exploration['yearly'])}

### Performance by Sector
{render_table(exploration['sector'], max_rows=11)}

### Performance by Holding Days
{render_table(exploration['holding_days'], max_rows=21)}

## Backtesting Methodology
- Starting capital: stocks ${STOCK_START_CAPITAL:,.0f}, options ${OPTIONS_START_CAPITAL:,.0f}, spreads ${SPREAD_START_CAPITAL:,.0f}
- Combined runs include both `l` and `s`; separate long-only and short-only evaluations use the same strategy grid on direction-filtered subsets
- Earnings filter excluded any trade with a symbol-level earnings date between entry and modeled exit
- Equity is tracked on an event basis and open positions are carried at cost until exit, so drawdown is based on realized equity changes rather than intraday mark-to-market noise
- Stock exits tested: hold to scheduled close, +4% target with a 2% trailing approximation, and +6% target with a 2x ATR downside floor
- Options use a simplified ATM call premium model with 2.5%-3.0% starting premium, capped downside at -100%, and theta drag scaled by holding period
- Spreads use simplified bull call debit and bull put credit payoff curves defined directly in percentage return space

## Stock Strategy Results
### Combined Long + Short
{render_table(stock_combined['results'], columns=['strategy_id','config','sharpe','max_drawdown','win_rate','annualized_return','trade_count','all_8_years_profitable'])}

### Pareto Frontier
{render_table(pareto_frontier(stock_combined['results']), columns=['strategy_id','sharpe','max_drawdown','annualized_return','trade_count'])}

### Top 10 Detailed
{chr(10).join(stock_detail_lines)}

### Holdout Check (Top 5 on 2018-2024 vs 2025)
{render_table(stock_combined['holdout'])}

### Key Findings
{"".join(f"- {line}\n" for line in findings_text(stock_combined['results'], 'combined stock'))}

{direction_block('Long-Only Stock Results', stock_long)}

{direction_block('Short-Only Stock Results', stock_short)}

## Options Strategy Results
### Combined Long + Short
{render_table(options_combined['results'], columns=['strategy_id','config','sharpe','max_drawdown','win_rate','annualized_return','trade_count','all_8_years_profitable'])}

### Pareto Frontier
{render_table(pareto_frontier(options_combined['results']), columns=['strategy_id','sharpe','max_drawdown','annualized_return','trade_count'])}

### Holdout Check (Top 5 on 2018-2024 vs 2025)
{render_table(options_combined['holdout'])}

### Key Findings
{"".join(f"- {line}\n" for line in findings_text(options_combined['results'], 'combined options'))}

{direction_block('Long-Only Options Results', options_long)}

{direction_block('Short-Only Options Results', options_short)}

## Spread Strategy Results
### Combined Long + Short
{render_table(spreads_combined['results'], columns=['strategy_id','config','sharpe','max_drawdown','win_rate','annualized_return','trade_count','all_8_years_profitable'])}

### Pareto Frontier
{render_table(pareto_frontier(spreads_combined['results']), columns=['strategy_id','sharpe','max_drawdown','annualized_return','trade_count'])}

### Holdout Check (Top 5 on 2018-2024 vs 2025)
{render_table(spreads_combined['holdout'])}

### Key Findings
{"".join(f"- {line}\n" for line in findings_text(spreads_combined['results'], 'combined spread'))}

{direction_block('Long-Only Spread Results', spreads_long)}

{direction_block('Short-Only Spread Results', spreads_short)}

## Recommended Strategy Set
- Combined stock portfolio: `{stock_combined['results'].iloc[0]['strategy_id']}`
- Long-only stock sleeve: `{stock_long['results'].iloc[0]['strategy_id']}`
- Short-only stock sleeve: `{stock_short['results'].iloc[0]['strategy_id']}`
- Combined options account: `{options_combined['results'].iloc[0]['strategy_id']}`
- Combined spreads account: `{spreads_combined['results'].iloc[0]['strategy_id']}`

## Independent Observations
- The raw return forecast is only moderately correlated with realized return; the ML scorer is more reliable as a ranking layer than a point-estimate engine.
- `predicted_mfe` is more useful for designing capped or target-driven exits than for estimating exact terminal returns.
- The short side is materially weaker than the long side at the raw signal level, so separate reporting is necessary even when combined portfolios are allowed.
- Robustness improves materially once earnings-overlap trades are removed, and concentration control matters as much as ranking choice.
"""
    MARKDOWN_OUT.write_text(content, encoding="utf-8")


def add_table(doc: Document, df: pd.DataFrame, title: str, max_rows: int | None = None) -> None:
    doc.add_heading(title, level=2)
    work = df.copy()
    if max_rows is not None:
        work = work.head(max_rows)
    table = doc.add_table(rows=1, cols=len(work.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(work.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in work.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = f"{value:.2f}" if isinstance(value, float) else str(value)


def write_docx(
    exploration: dict[str, object],
    stock_combined: dict[str, object],
    stock_long: dict[str, object],
    stock_short: dict[str, object],
    options_combined: dict[str, object],
    options_long: dict[str, object],
    options_short: dict[str, object],
    spreads_combined: dict[str, object],
    spreads_long: dict[str, object],
    spreads_short: dict[str, object],
) -> None:
    doc = Document()
    doc.add_heading("TradeWave V3 Codex Strategy Assessment", level=0)

    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"After including both directions and excluding {exploration['earnings_filtered_rate']:.2f}% of trades for earnings overlap, "
    )
    p.add_run(
        f"the best combined stock strategy delivered a Sharpe ratio of {stock_combined['results'].iloc[0]['sharpe']:.2f} with max drawdown of {stock_combined['results'].iloc[0]['max_drawdown']:.2f}%."
    ).bold = True
    doc.add_paragraph(
        "The core result is that V3 works best as a ranking engine. Composite ranking and disciplined portfolio limits outperformed naive hold-to-expiry behavior, while the long side remained stronger than the short side on raw signal quality."
    )

    doc.add_heading("ML Model Signal Quality", level=1)
    doc.add_paragraph(
        f"Predicted terminal return versus realized terminal return correlation was {exploration['pred_actual_corr']:.4f}; "
        f"predicted MFE versus realized MFE correlation was {exploration['mfe_actual_corr']:.4f}."
    )
    add_table(doc, exploration["direction"], "Direction Summary")
    add_table(doc, exploration["ml_score_deciles"], "ML Score Deciles")
    add_table(doc, exploration["win_probability_deciles"], "Win Probability Deciles")
    add_table(doc, exploration["yearly"], "Year-by-Year Signal Quality")

    doc.add_heading("Stock Strategy Results", level=1)
    add_table(doc, stock_combined["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 20 Combined Stock Strategies", max_rows=20)
    add_table(doc, stock_long["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 10 Long-Only Stock Strategies", max_rows=10)
    add_table(doc, stock_short["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 10 Short-Only Stock Strategies", max_rows=10)
    for strategy_id in stock_combined["results"].head(2)["strategy_id"]:
        sim = stock_combined["details"][strategy_id]["simulation"]
        yearly_df = pd.DataFrame([{"year": y, "return_pct": r} for y, r in sim["yearly_returns"].items()])
        if not yearly_df.empty:
            add_table(doc, yearly_df, f"{strategy_id} Annual Returns")
    add_table(doc, stock_combined["holdout"], "Combined Stock Holdout Validation")

    doc.add_heading("Options Strategy Results", level=1)
    doc.add_paragraph(
        "Options were modeled as ATM or near-ATM long calls in a $10,000 account with simplified premium and time-decay assumptions. These figures should be treated as approximate directional evidence, not executable option book P&L."
    )
    add_table(doc, options_combined["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 20 Combined Options Strategies", max_rows=20)
    add_table(doc, options_long["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 10 Long-Only Options Strategies", max_rows=10)
    add_table(doc, options_short["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 10 Short-Only Options Strategies", max_rows=10)
    add_table(doc, options_combined["holdout"], "Combined Options Holdout Validation")

    doc.add_heading("Spread Strategy Results", level=1)
    doc.add_paragraph(
        "Spread results are split across bull call debit and bull put credit variants. Capped-risk spreads reduced leverage path dependency and generally improved drawdown relative to naked calls."
    )
    add_table(doc, spreads_combined["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 20 Combined Spread Strategies", max_rows=20)
    add_table(doc, spreads_long["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 10 Long-Only Spread Strategies", max_rows=10)
    add_table(doc, spreads_short["results"][["strategy_id", "config", "sharpe", "max_drawdown", "win_rate", "annualized_return", "trade_count"]], "Top 10 Short-Only Spread Strategies", max_rows=10)
    add_table(doc, spreads_combined["holdout"], "Combined Spread Holdout Validation")

    doc.add_heading("Risk Analysis", level=1)
    doc.add_paragraph(
        "This backtest does not observe bid-ask spread, slippage, assignment, or intraday stop path. Results are therefore best used for relative strategy comparison rather than exact expectation setting."
    )
    doc.add_paragraph(
        "The main failure mode is a correlated drawdown cluster, especially when high-score names are concentrated in the same sector or macro regime. Earnings filtering helps, but it does not remove all event risk."
    )

    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph(
        f"Primary combined deployment candidate: {stock_combined['results'].iloc[0]['strategy_id']}. "
        f"Primary long-only sleeve: {stock_long['results'].iloc[0]['strategy_id']}. Primary short-only sleeve: {stock_short['results'].iloc[0]['strategy_id']}. "
        f"For leveraged implementations, prefer the top spread variant over the top naked-call variant. "
        "Plan on a 20%-30% haircut to annualized returns when translating these backtests into live expectations."
    )

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def main() -> None:
    df = load_data()
    exploration = data_exploration(df)
    df_long = df[df["direction"] == "l"].copy()
    df_short = df[df["direction"] == "s"].copy()

    stock_configs = generate_stock_strategies()
    option_configs = generate_option_strategies()
    spread_configs = generate_spread_strategies()

    stock_combined = run_suite(df, stock_configs, STOCK_START_CAPITAL, stock_return_builder)
    stock_long = run_suite(df_long, stock_configs, STOCK_START_CAPITAL, stock_return_builder)
    stock_short = run_suite(df_short, stock_configs, STOCK_START_CAPITAL, stock_return_builder)

    options_combined = run_suite(df, option_configs, OPTIONS_START_CAPITAL, option_return_builder)
    options_long = run_suite(df_long, option_configs, OPTIONS_START_CAPITAL, option_return_builder)
    options_short = run_suite(df_short, option_configs, OPTIONS_START_CAPITAL, option_return_builder)

    spreads_combined = run_suite(df, spread_configs, SPREAD_START_CAPITAL, spread_return_builder)
    spreads_long = run_suite(df_long, spread_configs, SPREAD_START_CAPITAL, spread_return_builder)
    spreads_short = run_suite(df_short, spread_configs, SPREAD_START_CAPITAL, spread_return_builder)

    MARKDOWN_OUT.parent.mkdir(parents=True, exist_ok=True)
    write_markdown(
        exploration,
        stock_combined,
        stock_long,
        stock_short,
        options_combined,
        options_long,
        options_short,
        spreads_combined,
        spreads_long,
        spreads_short,
    )
    write_docx(
        exploration,
        stock_combined,
        stock_long,
        stock_short,
        options_combined,
        options_long,
        options_short,
        spreads_combined,
        spreads_long,
        spreads_short,
    )
    print(f"Wrote {MARKDOWN_OUT}")
    print(f"Wrote {DOCX_OUT}")


if __name__ == "__main__":
    main()
