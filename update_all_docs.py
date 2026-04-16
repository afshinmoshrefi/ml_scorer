"""Update Options Playbook with spread findings and add return projections appendix."""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def fix_styles(doc):
    seen = set()
    remove = []
    for child in doc.styles.element:
        sid = child.get(qn("w:styleId"))
        if sid and sid in seen:
            remove.append(child)
        elif sid:
            seen.add(sid)
    for elem in remove:
        doc.styles.element.remove(elem)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pStyle")):
        pPr.remove(existing)
    pPr.insert(0, pPr.makeelement(qn("w:pStyle"), {qn("w:val"): f"Heading{level}"}))
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r, rd in enumerate(rows):
        for c, v in enumerate(rd):
            t.rows[r + 1].cells[c].text = str(v)
    return t


# ============================================================
# UPDATE OPTIONS PLAYBOOK with Spread Strategy Findings
# ============================================================

doc = Document("docs/TradeWave_Options_Strategy_Playbook_V2.docx")
fix_styles(doc)

doc.add_page_break()
add_heading(doc, "11. Spread Strategy Backtest Results", level=1)

doc.add_paragraph(
    "After the initial 116-strategy backtest revealed that naked long options "
    "struggle against theta decay, a second-generation backtester was built to test "
    "multi-leg spread strategies. These structures fundamentally change the theta "
    "dynamics by either reducing cost (debit spreads) or making theta work in the "
    "trader's favor (credit spreads)."
)

# --- 11.1 ---
add_heading(doc, "11.1 Structures Tested", level=2)

doc.add_paragraph(
    "Three option structures were tested against the same 8-year walk-forward data:"
)

add_table(doc,
    ["Structure", "Mechanics", "Theta Effect", "Max Loss", "Count"],
    [
        ["Bull Call Debit Spread", "Buy ATM call, sell OTM call at spread width above entry",
         "Partially offset (short leg decays too)", "Premium paid (35-40% of spread)", "38"],
        ["Bull Put Credit Spread", "Sell OTM put, buy further OTM put for protection",
         "Works FOR you (premium decays in your favor)", "Spread width minus credit", "22"],
        ["Deep ITM Single Leg", "Buy D60 call with X2 expiry (benchmark from V1)",
         "Minimal (low time value)", "Premium paid", "5"],
    ],
)

# --- 11.2 ---
add_heading(doc, "11.2 Results: Spreads vs Single-Leg Options", level=2)

add_table(doc,
    ["Structure", "Avg Sharpe", "Avg DD", "Avg WR", "All Profitable", "8yr+"],
    [
        ["Debit Spreads", "3.00", "54%", "70%", "38/38 (100%)", "37/38"],
        ["Credit Spreads", "2.15", "51%", "87%", "22/22 (100%)", "11/22"],
        ["Deep ITM (single)", "2.25", "41%", "71%", "5/5 (100%)", "3/5"],
        ["V1 single-leg (116 strats)", "0.55", "68%", "57%", "73/116 (63%)", "1/116"],
    ],
)

doc.add_paragraph(
    "Debit spreads improved Sharpe from 0.55 (V1 single-leg average) to 3.00 -- a 5.5x "
    "improvement. Every single debit spread strategy was profitable. 37 of 38 were "
    "profitable in all 8 years. This is the most dramatic finding in the entire "
    "options backtest program."
)

doc.add_paragraph(
    "Credit spreads achieved 2.15 Sharpe with 87% win rate. While not as high-Sharpe "
    "as debit spreads, they provide a psychologically comfortable trading experience "
    "with nearly 9 out of 10 trades winning."
)

# --- 11.3 ---
add_heading(doc, "11.3 Why Spreads Work", level=2)

add_heading(doc, "Debit Spreads: Lower Breakeven", level=3)
doc.add_paragraph(
    "A naked ATM call on a $100 stock costs ~$2.80 (20 DTE, 25% IV). The stock must "
    "rise 2.8% just to break even. A 3% bull call spread ($100/$103) costs ~$1.20 "
    "(40% of the $3 spread width). Breakeven is only $101.20 -- a 1.2% move. "
    "The ML model predicts 82% of ML_85 trades exceed 0% and ~60% exceed 3%. "
    "At a 1.2% breakeven, roughly 70% of trades are profitable."
)
doc.add_paragraph(
    "The short leg also partially offsets theta. While the long call decays, the "
    "short call decays too (in the trader's favor). The net theta is much lower "
    "than a naked call."
)

add_heading(doc, "Credit Spreads: Theta Works For You", level=3)
doc.add_paragraph(
    "A bull put credit spread profits when the stock stays above the short put strike. "
    "With the ML model predicting upward moves with 82-85% accuracy, the stock staying "
    "above a 2-3% OTM put is approximately 89% likely. You collect credit upfront and "
    "theta works in your favor -- every passing day, the premium you sold decays toward zero. "
    "This completely inverts the theta problem of long options."
)

# --- 11.4 ---
add_heading(doc, "11.4 Top Spread Strategies", level=2)

add_heading(doc, "Debit Spread #7 -- Best Risk-Adjusted", level=3)
doc.add_paragraph(
    "Configuration: WP ranking, T90, 3% spread width, XN expiry, Quarter Kelly, 3 positions, IL IV filter."
)
add_table(doc,
    ["Sharpe", "Max DD", "WR", "Trades", "EV/Trade", "PF", "Yr+"],
    [["3.69", "45.8%", "72.0%", "328", "+75% of premium", "1.70", "8"]],
)
doc.add_paragraph(
    "Year-by-year: 2018: 685%, 2019: 121%, 2020: 2059%, 2021: 2234%, "
    "2022: 1065%, 2023: 4982%, 2024: 725%, 2025: 1345%. "
    "Profitable all 8 years. Average winning trade returns +137% of premium paid. "
    "Average losing trade loses -84% of premium. With 72% win rate, the expected "
    "value is +75% per trade."
)
doc.add_paragraph(
    "Example trade: AAPL at $175. Buy $175 call, sell $180.25 call (3% spread). "
    "Cost: ~$2.10 ($1.20 per contract x $5.25 spread width x 40% cost ratio). "
    "If AAPL reaches $180.25+: profit $3.15 (+150%). If AAPL stays flat: lose $2.10 (-100%). "
    "At 72% WR, expected profit per $1 risked: $0.75."
)

add_heading(doc, "Credit Spread #47 -- Highest Win Rate", level=3)
doc.add_paragraph(
    "Configuration: CW ranking, T85, 2% OTM short put, 3% spread width, "
    "XS expiry, Quarter Kelly, 3 positions, IL IV filter."
)
add_table(doc,
    ["Sharpe", "Max DD", "WR", "Trades", "EV/Trade", "PF", "Yr+"],
    [["3.17", "38.5%", "89.2%", "323", "+33% of collateral", "7.35", "8"]],
)
doc.add_paragraph(
    "Year-by-year: 2018: 181%, 2019: 19%, 2020: 157%, 2021: 343%, "
    "2022: 203%, 2023: 264%, 2024: 400%, 2025: 387%. "
    "Profitable all 8 years. 89% of trades are full wins (keep entire credit). "
    "Average win: +48% of collateral. Average loss: -94% of collateral. "
    "But with 89% win rate, expected value is +33% per trade on collateral."
)
doc.add_paragraph(
    "Example trade: MSFT at $400. Sell $392 put (2% OTM), buy $380 put (protection). "
    "Collect ~$3.56 credit on $12 spread (30% credit ratio). Max risk: $8.44. "
    "If MSFT stays above $392 (89% of the time): keep $3.56 (+42% on collateral). "
    "If MSFT drops below $380: lose $8.44 (-100% on collateral)."
)

# --- 11.5 ---
add_heading(doc, "11.5 Per-Trade Economics", level=2)

add_table(doc,
    ["Strategy", "Win Rate", "Avg Win", "Avg Loss", "EV/Trade", "Trades/Yr"],
    [
        ["Stock #23 (EP trail)", "56%", "+9.1%", "-3.2%", "+4.0%", "87"],
        ["Stock #6 (MFE trail)", "85%", "+7.2%", "-7.4%", "+5.1%", "55"],
        ["Debit Spread #7", "72%", "+137%*", "-84%*", "+75%*", "41"],
        ["Credit Spread #47", "89%", "+48%*", "-94%*", "+33%*", "40"],
    ],
)
doc.add_paragraph("* = percentage of premium paid (debit) or collateral (credit), not of account.")

# --- 11.6 ---
add_heading(doc, "11.6 Return Projections (Non-Compounding)", level=2)
doc.add_paragraph(
    "These projections use simple (non-compounding) math: "
    "annual return = EV per trade x trades per year x allocation percentage. "
    "No reinvestment of profits. This is the floor estimate."
)

add_heading(doc, "Stock Strategy #23 ($100K-$5M accounts)", level=3)
add_table(doc,
    ["Account Size", "Annual Return ($)", "Annual Return (%)", "Worst Year"],
    [
        ["$100,000", "$59,607", "60%", "$26,374 (2019)"],
        ["$500,000", "$298,037", "60%", "$131,870 (2019)"],
        ["$1,000,000", "$596,075", "60%", "$263,740 (2019)"],
        ["$5,000,000", "$2,980,373", "60%", "$1,318,700 (2019)"],
    ],
)

add_heading(doc, "Debit Spread #7 ($10K-$100K accounts)", level=3)
add_table(doc,
    ["Account Size", "Annual Return ($)", "Annual Return (%)", "Worst Year"],
    [
        ["$10,000", "$30,816", "308%", "$7,622 (2019)"],
        ["$25,000", "$77,040", "308%", "$19,055 (2019)"],
        ["$50,000", "$154,079", "308%", "$38,110 (2019)"],
    ],
)
doc.add_paragraph(
    "Note: The 308% return on capital reflects the leverage of options. "
    "Only 10% of capital is deployed per trade (max 50% total), so the actual "
    "dollars at risk at any time are $1K-$5K on a $10K account. The remaining "
    "capital is cash reserve. Losing the full premium on 3 trades simultaneously "
    "would cost ~$3K (30% of a $10K account)."
)

add_heading(doc, "Credit Spread #47 ($10K-$50K accounts)", level=3)
add_table(doc,
    ["Account Size", "Annual Return ($)", "Annual Return (%)", "Worst Year"],
    [
        ["$10,000", "$13,300", "133%", "$1,396 (2019)"],
        ["$25,000", "$33,250", "133%", "$3,490 (2019)"],
        ["$50,000", "$66,499", "133%", "$6,980 (2019)"],
    ],
)

add_heading(doc, "Combined Portfolio", level=3)
add_table(doc,
    ["Scenario", "Stock ($500K)", "Options ($25K)", "Total ($525K)", "Return"],
    [
        ["Backtest", "$298,037", "$77,040", "$375,077", "71%"],
        ["Conservative (-30%)", "$208,626", "$53,928", "$262,554", "50%"],
        ["Pessimistic (-50%)", "$149,019", "$38,520", "$187,538", "36%"],
    ],
)
doc.add_paragraph(
    "Even at the pessimistic 50% haircut (accounting for execution issues, "
    "model decay, timing slippage, and bid-ask spreads wider than modeled), "
    "the combined portfolio projects $188K/yr on $525K (36% annual return). "
    "The stock component provides steady base returns; the options component "
    "adds outsized returns on a small capital allocation."
)

# --- 11.7 ---
add_heading(doc, "11.7 Risk Metrics", level=2)

add_table(doc,
    ["Metric", "Stock #23", "Stock #6", "DS #7", "CS #47"],
    [
        ["Worst single trade", "-3.2%", "-36.9%", "-100%*", "-100%*"],
        ["5th percentile trade", "-3.2%", "-9.3%", "-100%*", "-100%*"],
        ["Worst month (sum of trades)", "-13.2%", "-46.7%", "-300%*", "-300%*"],
        ["Max losing streak", "11 trades", "6 trades", "5 trades", "5 trades"],
        ["Streaks of 3+ losses", "48x in 8yr", "3x in 8yr", "6x in 8yr", "3x in 8yr"],
    ],
)
doc.add_paragraph(
    "* = percentage of premium/collateral, not of account. A -100% option trade "
    "on 10% allocation = -10% of account. Three simultaneous -100% trades = -30% of account. "
    "Stock #23's worst single trade is always -3.2% (the EP trailing stop cap). "
    "Stock #6 can have larger individual losses (-36.9%) but they are rare (15% of trades)."
)

# --- 11.8 ---
add_heading(doc, "11.8 Recommendation Update", level=2)
doc.add_paragraph(
    "The spread backtest changes the options recommendation entirely. "
    "The original playbook tested 116 single-leg strategies and found only 1 "
    "profitable all 8 years. The spread backtester found 48 out of 60 profitable "
    "all 8 years. Spreads are not an incremental improvement -- they are a "
    "fundamentally different risk profile."
)
doc.add_paragraph(
    "Updated recommendation for the $10K options account: "
    "Primary: Debit Spread #7 (bull call spread, 3% width, T90, XN expiry, SK sizing). "
    "Secondary: Credit Spread #47 (bull put spread, 2% OTM, 3% width, XS expiry, SK sizing). "
    "Both can run simultaneously on the same account since they use different capital "
    "(debit spreads use premium; credit spreads use collateral/margin)."
)
doc.add_paragraph(
    "Key rules: (1) Never buy naked OTM calls on seasonal patterns. "
    "(2) Debit spreads: keep width at 3-5%, use XN or longer expiry. "
    "(3) Credit spreads: short put 2-3% OTM, XS expiry for fastest theta collection. "
    "(4) Both: Quarter Kelly sizing, max 10% of capital per trade, 50% total deployment cap."
)

# --- 11.9 ---
add_heading(doc, "11.9 Output Data Files", level=2)
doc.add_paragraph(
    "Spread backtest output in results/backtest_spreads/: "
    "summary.csv (65 rows), trades.csv (23,852 rows), equity.csv (130,715 rows)."
)

doc.save("docs/TradeWave_Options_Strategy_Playbook_V2.docx")
print("Options playbook updated with Section 11 (Spread Results + Return Projections).")


# ============================================================
# UPDATE STOCK PLAYBOOK with Return Projections
# ============================================================

doc2 = Document("docs/TradeWave_Stock_Strategy_Playbook_V2.docx")
fix_styles(doc2)

doc2.add_page_break()
add_heading(doc2, "12. Return Projections", level=1)

doc2.add_paragraph(
    "Non-compounding return projections based on per-trade economics from the 8-year "
    "walk-forward backtest. These are floor estimates -- compounding would increase returns "
    "but is subject to practical liquidity limits at larger account sizes."
)

add_heading(doc2, "12.1 Per-Trade Economics", level=2)
add_table(doc2,
    ["Strategy", "Win Rate", "Avg Win", "Avg Loss", "EV/Trade", "Trades/Yr"],
    [
        ["#23 (EP/T90/SK/C1)", "56.2%", "+9.1%", "-3.2%", "+4.0%", "87"],
        ["#6 (EM/T90/SK/C1)", "85.3%", "+7.2%", "-7.4%", "+5.1%", "55"],
    ],
)
doc2.add_paragraph(
    "Strategy #23 has a bounded loss profile: 85% of all losses are exactly -3.2% "
    "(the 3% EP trailing stop plus 0.2% slippage). Winners average +9.1% with a "
    "fat right tail (best trade: +107%). This asymmetry drives the 4.0% EV per trade."
)
doc2.add_paragraph(
    "Strategy #6 has a high win rate profile: only 14.7% of trades lose, but losses "
    "average -7.4% (no loss cap like EP). The higher EV per trade (+5.1%) comes from "
    "the MFE trailing stop capturing more upside on winners."
)

add_heading(doc2, "12.2 Dollar Projections", level=2)
add_table(doc2,
    ["Account Size", "#23 Annual ($)", "#23 Annual (%)", "#6 Annual ($)", "#6 Annual (%)"],
    [
        ["$100,000", "$59,607", "60%", "$47,570", "48%"],
        ["$250,000", "$149,019", "60%", "$118,925", "48%"],
        ["$500,000", "$298,037", "60%", "$237,849", "48%"],
        ["$1,000,000", "$596,075", "60%", "$475,698", "48%"],
    ],
)
doc2.add_paragraph(
    "Computed as: EV/trade x trades/year x 17% allocation (Quarter Kelly with 3 positions). "
    "Non-compounding. These estimates scale linearly with account size up to the point "
    "where position sizes impact market liquidity (~$5-10M for most S&P 500 stocks)."
)

add_heading(doc2, "12.3 Year-by-Year (Strategy #23, $100K)", level=2)
add_table(doc2,
    ["Year", "Trades", "EV/Trade", "Annual $", "Annual %"],
    [
        ["2018", "80", "+3.5%", "+$47,536", "+48%"],
        ["2019", "11", "+14.1%", "+$26,374", "+26%"],
        ["2020", "89", "+4.2%", "+$63,800", "+64%"],
        ["2021", "91", "+4.4%", "+$68,139", "+68%"],
        ["2022", "139", "+3.2%", "+$74,436", "+74%"],
        ["2023", "102", "+4.6%", "+$79,682", "+80%"],
        ["2024", "90", "+4.0%", "+$61,574", "+62%"],
        ["2025", "96", "+3.4%", "+$55,318", "+55%"],
        ["Average", "87", "+4.0%", "+$59,607", "+60%"],
    ],
)
doc2.add_paragraph(
    "The worst year (2019) still produced +$26K (+26%) despite having only 11 trades. "
    "The best year (2023) produced +$80K (+80%). The narrow range (26-80%) demonstrates "
    "the consistency of the cash machine approach."
)

add_heading(doc2, "12.4 Degradation Scenarios", level=2)
add_table(doc2,
    ["Scenario", "#23 Return", "#6 Return", "Assumption"],
    [
        ["Backtest", "60%", "48%", "Per-trade economics as measured"],
        ["Conservative (-30%)", "42%", "33%", "Execution gaps, wider spreads, timing"],
        ["Pessimistic (-50%)", "30%", "24%", "Model decay, regime change, worst case"],
    ],
)
doc2.add_paragraph(
    "Even at the pessimistic 50% haircut, Strategy #23 projects 30% annual simple return "
    "on the stock account. This represents a realistic floor for the first year of live "
    "trading before the model has been validated in production."
)

doc2.save("docs/TradeWave_Stock_Strategy_Playbook_V2.docx")
print("Stock playbook updated with Section 12 (Return Projections).")
