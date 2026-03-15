"""Generate ML Pattern Scorer V2 Production Readiness Report as .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def add_table(headers, rows):
    """Add a table using a native Word built-in style."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if c_idx == 0:
                run.bold = True

    doc.add_paragraph()  # spacing after table


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('ML Pattern Scorer V2', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Deployment Readiness Analysis')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('March 2026  |  10-30 Day Tier  |  SR + MFE Ensemble Models')
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'The V2 ML Pattern Scorer is a 3-model ensemble (LightGBM + XGBoost + CatBoost) '
    'trained on 34.7 million historical samples across 475 S&P 500 stocks and 26 years '
    '(2000-2025). It scores seasonal stock pattern opportunities by predicting both '
    'actual return (SR model) and maximum favorable excursion (MFE model).'
)
doc.add_paragraph(
    'Eight years of true out-of-sample walk-forward validation (2018-2025) demonstrate '
    'that filtering to the model\'s top 30% of scored opportunities (ML_70 threshold) '
    'consistently delivers cohort-level win rates of 78-86% and cohort-level Sharpe ratios of 7-12, '
    'compared to baseline win rates of 61-82% and Sharpe ratios of 2-9. The ML_70 '
    'cohort has never had a negative average return in any validation year.'
)
doc.add_paragraph(
    'The model is ready for production deployment in the TradeWave UI (opportunity '
    'ranking and scoring) and ready for controlled live testing in automated trading '
    '(shadow mode or small capital with safeguards). Portfolio construction, trade '
    'selection policy, and live monitoring remain to be built and validated.'
)

# ============================================================
# 1. FILTERING POWER
# ============================================================
doc.add_heading('1. Model Filtering Power', level=1)
doc.add_paragraph(
    'The core value of the model is opportunity filtering: separating the best seasonal '
    'patterns from the rest. The table below shows ML_70 (top 30%) performance vs. '
    'unfiltered baseline across all 8 walk-forward validation years.'
)
doc.add_paragraph(
    'Important: all win rates, average returns, and Sharpe ratios in this report are '
    'cohort statistics computed across the full universe of scored opportunities in each '
    'validation year. They represent equal-weight averages over all qualifying patterns, '
    'not realized portfolio returns from a specific trading strategy. Actual portfolio '
    'performance will depend on trade selection policy, position sizing, execution, and '
    'the number of concurrent positions held.'
)

add_table(
    ['Year', 'Base WR', 'ML_70 WR', 'WR Lift', 'Base Sharpe', 'ML_70 Sharpe', 'Sharpe Lift'],
    [
        ['2018', '71.7%', '77.8%', '+6.1', '4.91', '7.04', '+43%'],
        ['2019', '64.2%', '77.8%', '+13.6', '3.26', '6.93', '+113%'],
        ['2020', '62.4%', '78.2%', '+15.8', '2.13', '8.32', '+291%'],
        ['2021', '68.9%', '78.8%', '+9.9', '5.34', '8.68', '+63%'],
        ['2022', '81.6%', '86.0%', '+4.4', '9.36', '12.31', '+32%'],
        ['2023', '65.0%', '79.7%', '+14.7', '4.08', '9.00', '+121%'],
        ['2024', '65.9%', '77.8%', '+11.9', '4.44', '8.21', '+85%'],
        ['2025', '61.5%', '78.1%', '+16.6', '2.33', '7.45', '+220%'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Key finding: ')
r.bold = True
p.add_run(
    'The model adds the most value in hard markets. When baseline WR is low '
    '(2020: 62.4%, 2025: 61.5%), the model lifts WR by 15-17 points and Sharpe by '
    '220-291%. When the market is already easy (2022: 81.6% base WR), the model adds '
    'only 4 points. This is the right behavior -- the gatekeeper matters most when '
    'conditions are worst.'
)

p2 = doc.add_paragraph()
r2 = p2.add_run('ML_70 cohort win rate never drops below 77.8%. ')
r2.bold = True
p2.add_run(
    'That is the floor across 8 diverse years including COVID (2020), a bear market '
    '(2022), and rate hiking cycles. This means that in every validation year, the '
    'average opportunity in the model\'s top 30% was profitable at least 77.8% of the '
    'time -- a consistent edge across regimes.'
)

doc.add_heading('Reading the Sharpe Ratios', level=2)
doc.add_paragraph(
    'The Sharpe ratios in this report (7-12 at ML_70) are cohort-level statistics, '
    'not portfolio-level Sharpe ratios. The distinction matters:'
)
doc.add_paragraph(
    'Cohort Sharpe = mean(all opportunity returns in cohort) / std(all opportunity '
    'returns in cohort). This is computed across hundreds of thousands of individual '
    'pattern outcomes per validation year. The massive sample sizes produce a stable '
    'mean and compress the ratio upward. This measures the quality of the scoring '
    'filter: how well the model separates good opportunities from bad ones.'
)
doc.add_paragraph(
    'Portfolio Sharpe is what a finance professional typically means by "Sharpe ratio." '
    'It measures annualized risk-adjusted return of an actual trading account: '
    'annualized_return / annualized_volatility. A hedge fund Sharpe of 2.0 is '
    'considered exceptional. The cohort Sharpe of 7-12 is not comparable to this number.'
)
doc.add_paragraph(
    'Why the difference is so large: a real portfolio holds 3-4 concurrent positions, '
    'not 200,000. Individual position variance dominates. The portfolio also faces '
    'execution costs, bid-ask spreads, timing risk, and correlation between concurrent '
    'positions. All of these increase realized volatility and reduce the ratio.'
)
p3 = doc.add_paragraph()
r3 = p3.add_run('Estimated realistic portfolio Sharpe: ')
r3.bold = True
p3.add_run(
    'A well-constructed portfolio using this scorer with 3-4 concurrent positions, '
    'sensible trade selection, and proper risk management would likely produce an '
    'annualized Sharpe in the 1.0-3.0 range. This is still a strong result -- most '
    'active equity strategies target 0.5-1.5. The cohort Sharpe of 7-12 should be '
    'read as "the filter is very good at separating winners from losers," not as '
    '"this strategy produces a Sharpe of 10."'
)

# ============================================================
# 2. CALIBRATION
# ============================================================
doc.add_heading('2. Calibration Quality', level=1)

doc.add_heading('SR Win Probability Calibration', level=2)
doc.add_paragraph(
    'The SR model\'s predicted return maps smoothly and monotonically to observed win '
    'probability. Calibration built from 11 million walk-forward samples across 8 years.'
)

add_table(
    ['Prediction Bin', 'Pred Range', 'Win Prob', 'Avg Return', 'P(Hit Pred)'],
    [
        ['Bin 0 (lowest)', '< -0.56%', '51.1%', '-0.27%', '57.6%'],
        ['Bin 5', '~0.61-0.72%', '56.0%', '0.16%', '49.4%'],
        ['Bin 10', '~1.12-1.25%', '69.2%', '1.59%', '57.5%'],
        ['Bin 15', '~1.92-2.15%', '76.6%', '2.94%', '57.8%'],
        ['Bin 19 (highest)', '> 3.44%', '84.5%', '6.97%', '60.6%'],
    ]
)

doc.add_paragraph(
    'Win probability calibration is excellent: smooth, monotonic, and usable for '
    'production display. P(hit predicted return) is weaker (54-61%), meaning the '
    'model\'s magnitude estimates are noisy. Use predicted return as a ranking signal, '
    'not as a precise target.'
)

doc.add_heading('MFE Calibration', level=2)

add_table(
    ['Prediction Bin', 'Pred MFE Range', 'Win Prob', 'Avg MFE', 'P(Hit Pred MFE)'],
    [
        ['Bin 0 (lowest)', '0.57-2.64%', '57.2%', '2.58%', '46.9%'],
        ['Bin 5', '3.71-3.93%', '64.4%', '4.03%', '46.1%'],
        ['Bin 10', '4.82-5.07%', '67.2%', '5.14%', '45.6%'],
        ['Bin 15', '6.33-6.78%', '71.7%', '6.94%', '47.3%'],
        ['Bin 19 (highest)', '> 9.48%', '80.5%', '13.41%', '52.4%'],
    ]
)

doc.add_paragraph(
    'The MFE model correctly ranks opportunities by excursion potential (Spearman 0.37 '
    'rank correlation) but P(hit predicted MFE) is nearly flat at 46-52%. The model '
    'identifies WHICH opportunities have high upside, but the exact predicted MFE level '
    'should not be used as a literal profit target.'
)

# ============================================================
# 3. ENSEMBLE VALUE
# ============================================================
doc.add_heading('3. Ensemble Value', level=1)
doc.add_paragraph(
    'The 3-model ensemble provides critical insurance against individual model failure. '
    'No single model dominates across all years, and any one can fail in a given regime.'
)

add_table(
    ['Year', 'LightGBM', 'XGBoost', 'CatBoost', 'Ensemble'],
    [
        ['2018', '0.618', '0.601', '0.615', '0.614'],
        ['2019', '0.490 *', '0.592', '0.666', '0.637'],
        ['2020', '0.663', '0.641', '0.618', '0.651'],
        ['2021', '0.601', '0.603', '0.620', '0.611'],
        ['2022', '0.550', '0.569', '0.546', '0.564'],
        ['2023', '0.616', '0.630', '0.645', '0.644'],
        ['2024', '0.581', '0.630', '0.608', '0.623'],
        ['2025', '0.669', '0.667', '0.640', '0.671'],
    ]
)

p = doc.add_paragraph()
p.add_run('* ').bold = True
p.add_run(
    'In 2019, LightGBM produced essentially a null model (AUC 0.49, prediction '
    'std = 0.05). XGBoost and CatBoost rescued the ensemble. This single example '
    'justifies the 3-model architecture.'
)
doc.add_paragraph(
    'The ensemble never falls below AUC 0.564, while individual models hit 0.490. '
    'No single model is consistently best: LGB leads in 2025, CB saves 2019, XGB '
    'leads 2024. The ensemble provides stability.'
)

# ============================================================
# 4. CONCERNS AND RISKS
# ============================================================
doc.add_heading('4. Concerns and Risks', level=1)

doc.add_heading('A. The Model is a Weak Classifier', level=2)
doc.add_paragraph(
    'Average AUC of ~0.62 is above random (0.50) but below the "strong" threshold '
    '(0.70+). The model works because: (1) the base rate is already favorable -- '
    'seasonal patterns have inherent edge; (2) filtering top 30% amplifies a small '
    'signal across huge sample sizes; (3) the model only needs to be directionally '
    'right, not precisely calibrated. Do not expect the model to turn a bad pattern '
    'into a good one. It separates "good enough" from "don\'t bother."'
)

doc.add_heading('B. Training Instability', level=2)
doc.add_paragraph(
    'Optimal model complexity varies wildly: best iteration ranges from 2 (2019) to '
    '834 (2025). This means the signal-to-noise ratio changes drastically across '
    'market regimes. The ensemble masks this instability, but it indicates that the '
    'underlying seasonal signal is regime-dependent.'
)

doc.add_heading('C. Max Drawdown Risk', level=2)
doc.add_paragraph(
    'Even at ML_70+ thresholds, individual trades can produce large losses. For stock '
    'positions, the worst single-trade returns in historical validation are in the '
    '-20% to -40% range (e.g., earnings gaps, COVID-era crashes). For options, these '
    'same moves can result in total premium loss since options amplify underlying '
    'moves. Risk management (position sizing, stop losses, sector limits) is '
    'non-negotiable for automated trading.'
)

doc.add_heading('D. MFE Model Does Not Help Trade Selection', level=2)
doc.add_paragraph(
    'The MFE model identifies opportunities with high upside potential but should NOT '
    'be used for trade entry decisions. Comparison at ML_70:'
)

add_table(
    ['Year', 'SR ML_70 WR', 'MFE ML_70 WR', 'SR ML_70 Sharpe', 'MFE ML_70 Sharpe'],
    [
        ['2019', '77.8%', '64.4%', '6.93', '3.29'],
        ['2020', '78.2%', '72.7%', '8.32', '6.18'],
        ['2023', '79.7%', '73.3%', '9.00', '6.92'],
        ['2025', '78.1%', '68.2%', '7.45', '3.97'],
    ]
)

doc.add_paragraph(
    'SR filtering is dramatically better for trade selection. Use the SR model for '
    'go/no-go decisions. Use the MFE model only as supplementary information for '
    'profit target guidance.'
)

# ============================================================
# 5. TRADEWAVE UI RECOMMENDATIONS
# ============================================================
doc.add_heading('5. Recommendations for TradeWave UI', level=1)
doc.add_paragraph('The model is well-suited for UI integration. Recommended columns:')

bullets = [
    ('ML Score (0-100)', 'Predicted return percentile. Color code: green (70+), yellow (50-70), red (<50).'),
    ('Win Probability', 'From SR calibration table. Well-calibrated, monotonic. Display as percentage.'),
    ('Predicted Return', 'Raw SR model output. Useful for sorting. Display as percentage with one decimal.'),
    ('Upside Potential (MFE)', 'Raw MFE model output. Shows excursion potential. Display as percentage.'),
]
for title_text, desc in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Do NOT display in UI: ')
r.bold = True
p.add_run(
    'P(hit predicted return) and P(hit predicted MFE). These are useful for internal '
    'decision-making but confusing for users -- a "good" opportunity showing only 58% '
    'chance of hitting its predicted return is misleading without context.'
)

# ============================================================
# 6. AUTOMATED TRADING RECOMMENDATIONS
# ============================================================
doc.add_heading('6. Recommendations for Automated Trading', level=1)

doc.add_paragraph(
    'The scorer validates the first stage of an automated trading pipeline: opportunity '
    'scoring and filtering. The stages that remain unvalidated are trade selection policy '
    '(which 3-4 trades to pick from hundreds of daily candidates), portfolio construction, '
    'execution, and live monitoring. This section covers safeguards and preliminary strategy '
    'parameters for controlled live testing.'
)

doc.add_heading('The Trade Selection Gap', level=2)
doc.add_paragraph(
    'At ML_85, roughly 200,000 opportunities per year (~800 per day) pass the scoring '
    'threshold. The real system can only take 3-4 concurrent positions. The policy that '
    'selects which few names survive from this candidate pool is as important as the scorer '
    'itself. This selection policy (sorting by predicted return, sector diversification, '
    'correlation avoidance, position timing) needs to be treated as a second model that '
    'must be designed, tested, and validated separately.'
)

doc.add_heading('Required Safeguards', level=2)
safeguards = [
    ('Position sizing', 'No single position > 5-10% of capital ($500-1000 on $10K). With ~78% WR and options that can go to zero, small positions are essential.'),
    ('Max concurrent positions', '3-4 as planned. ML_85+ gives ~200K opportunities/year (~800/day). Need secondary sorting (by predicted return, sector diversification) to select best 3-4.'),
    ('Stop loss on options', '40-50% premium loss. Stock-level drawdowns of -20% to -40% get amplified in options and can result in total premium loss.'),
    ('Sector limits', 'Max 2 per sector. The model does not account for correlated sector risk.'),
    ('VIX circuit breaker', 'Do not score when VIX > 35. The model was trained with these samples removed and has no data for panic regimes.'),
]
for title_text, desc in safeguards:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Conservative Strategy at ML_85', level=2)
doc.add_paragraph(
    'For $10K starting capital with options, the ML_85 threshold provides the best '
    'balance of selectivity and opportunity volume:'
)

add_table(
    ['Year', 'ML_85 WR', 'ML_85 Avg Return', 'ML_85 Sharpe', 'Trades Available'],
    [
        ['2018', '80.2%', '3.34%', '8.11', '210K'],
        ['2019', '81.6%', '3.45%', '7.73', '213K'],
        ['2020', '81.8%', '7.06%', '10.67', '173K'],
        ['2021', '82.1%', '4.39%', '9.95', '212K'],
        ['2022', '89.2%', '7.78%', '14.46', '210K'],
        ['2023', '84.4%', '5.93%', '11.30', '212K'],
        ['2024', '78.3%', '4.62%', '8.72', '212K'],
        ['2025', '81.7%', '4.93%', '9.13', '209K'],
    ]
)

doc.add_paragraph(
    'ML_85 cohort statistics show ~82% average win rate, ~5.2% average return per '
    'trade, Sharpe ~10 across the filtered universe. These are not portfolio return '
    'projections. Actual results will depend heavily on which 40-80 trades per year '
    'the selection policy picks from the ~200K candidates, and on execution quality.'
)
doc.add_paragraph(
    'Note on options: the validation data measures stock-level close-to-close returns '
    'and maximum favorable excursion. Translating these to options P&L introduces '
    'additional variables (strike selection, implied volatility, theta decay, bid-ask '
    'spreads) that are not captured in the model. The scorer identifies favorable stock '
    'patterns; the options strategy layered on top requires its own validation.'
)

# ============================================================
# 7. MODEL ARCHITECTURE SUMMARY
# ============================================================
doc.add_heading('7. Model Architecture Summary', level=1)

arch_items = [
    ('Ensemble', '3 gradient boosting models (LightGBM + XGBoost + CatBoost), predictions averaged'),
    ('Two models per tier', 'SR (predicts actual return) and MFE (predicts max favorable excursion)'),
    ('Features', '59 features across 6 groups: pattern intrinsic, technical, market regime, context, calendar, interactions'),
    ('Training data', '34.7 million samples, 475 S&P 500 stocks, 26 years (2000-2025)'),
    ('Hyperparameters', 'Optuna-tuned (75 trials per target on 2M sample subset)'),
    ('Validation', '8-year expanding-window walk-forward (train 2000-Y, validate Y+1)'),
    ('VIX filter', 'Samples with VIX > 35 removed from training (~4.8% of data)'),
    ('Multi-tier', 'Architecture supports multiple day-range tiers (10-30, 31-60, etc.)'),
]
for title_text, desc in arch_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

# ============================================================
# 8. HOW THE MODEL ADAPTS TO MARKET CONDITIONS
# ============================================================
doc.add_heading('8. How the Model Adapts to Market Conditions', level=1)

doc.add_paragraph(
    'A common question is whether the model can adapt to different market environments -- '
    'bearish regimes, presidential election cycles, or directional trades (longs vs. shorts). '
    'The model does not use hard-coded rules for any of these. Instead, it learns the '
    'relationships between features and outcomes from 26 years of training data. This '
    'section explains how that works in practice.'
)

doc.add_heading('Presidential Election Cycle Awareness', level=2)
doc.add_paragraph(
    'The presidential election (PE) cycle is one of the strongest known seasonal effects '
    'in equity markets. Midterm years (PE year 2) tend to have weak summers, while '
    'pre-election years (PE year 3) tend to be strong. The model captures this through '
    'multiple features working together:'
)

pe_items = [
    ('cal_pe_year',
     'Tells the model which phase of the 4-year cycle we are in. '
     'The model learned from training data that patterns behave differently in each phase.'),
    ('pat_pe_match and pat_pe_deepest',
     'These features describe how a specific pattern performs when filtered to only PE-matching years. '
     'A pattern that is 30 years deep overall but only 4 years deep in PE-filtered history signals '
     'that the pattern does not have a strong track record in the current cycle phase.'),
    ('pat_pe_utilization',
     'Measures what fraction of available PE-cycle years the pattern passed. '
     'Low utilization means the pattern is inconsistent across election cycles.'),
    ('mkt_spx_seasonal_wr and mkt_spx_seasonal_ret',
     'These are computed from historical S&P 500 returns filtered by PE cycle year. '
     'They tell the model whether the broad market itself has a seasonal tendency '
     'to rise or fall during this calendar window in this type of election year.'),
    ('mkt_spx_seasonal_regime',
     'Classifies the SPX seasonal environment as bullish, neutral, or bearish based '
     'on the PE-filtered historical win rate. A bearish regime during midterm summer '
     'directly suppresses scores for long patterns.'),
]
for title_text, desc in pe_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Example: A Long Pattern That Fails in Midterm Years', level=2)
doc.add_paragraph(
    'Consider a long pattern on a stock like AAPL with a start date in June, holding 20 days. '
    'Suppose this pattern has strong overall history: 25 years deep, 85% win rate across all years, '
    'high Sharpe ratio. On the surface, this looks like an excellent opportunity.'
)
doc.add_paragraph(
    'However, when filtered to PE year 2 (midterm) years only, suppose the pattern is only 3 '
    'years deep out of 6 possible midterm years. That is a 50% win rate in midterm years '
    'specifically, compared to 85% overall. The model sees this divergence through several signals:'
)

example_items = [
    'cal_pe_year = 2, telling the model this is a midterm year.',
    'pat_pe_deepest = 3, showing shallow PE-filtered depth.',
    'pat_pe_utilization is low (~0.50), meaning the pattern passes in only half of available midterm years.',
    'mkt_spx_seasonal_regime is likely bearish, because SPX itself tends to decline in midterm summers.',
    'mkt_spx_dir_alignment = 0, because the long direction conflicts with the bearish SPX seasonal regime.',
]
for item in example_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_paragraph(
    'The model combines all of these signals and produces a lower predicted return, lower win '
    'probability, and lower ML score. The same pattern scored in a pre-election year (PE year 3) '
    'with strong PE depth would receive a significantly higher score. The model does not need '
    'an explicit "reject midterm longs" rule -- it learned this behavior from 26 years of outcomes.'
)

doc.add_heading('Long vs. Short Scoring', level=2)
doc.add_paragraph(
    'The model scores long and short patterns using the same 59 features and the same trained '
    'models. It does not have separate models for longs and shorts (this was tested and did not '
    'help -- halving the training data hurt more than direction specialization helped). Instead, '
    'direction is captured through features:'
)

dir_items = [
    ('pat_direction',
     'Binary feature: 1 for long, 0 for short. The model learned different return '
     'distributions for each direction across different market regimes.'),
    ('mkt_spx_dir_alignment',
     'Whether the pattern direction agrees with the SPX seasonal tendency. '
     'A short pattern in a bearish SPX seasonal regime gets alignment = 1 (favorable). '
     'A long pattern in the same regime gets alignment = 0 (unfavorable).'),
    ('pat_dir_x_mkt_trend',
     'Interaction feature that combines direction with broad market momentum. '
     'When SPY is trending down and the pattern is short, this feature is positive '
     '(supportive). When SPY is trending down and the pattern is long, this feature '
     'is negative (headwind).'),
    ('pat_dir_x_sector_trend',
     'Same concept applied to the stock\'s sector ETF trend. A short on an energy stock '
     'when XLE is falling gets a favorable signal.'),
]
for title_text, desc in dir_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

doc.add_paragraph(
    'In practice, this means that during a bearish midterm summer, long patterns will '
    'generally receive lower scores while short patterns with strong depth profiles will '
    'receive higher scores. The model naturally shifts its recommendations toward the '
    'direction that the current environment supports.'
)

doc.add_heading('Real-Time Market Regime Features', level=2)
doc.add_paragraph(
    'Beyond the calendar and PE cycle features, the model also uses 16 real-time market '
    'regime features that reflect current conditions at the moment of scoring. These include '
    'VIX level and term structure, yield curve slope, credit spreads, S&P 500 momentum and '
    'breadth, sector rotation, and federal funds rate. When bearish conditions materialize '
    '(rising VIX, inverting yield curve, widening credit spreads, deteriorating breadth), '
    'these features compound the penalty on long patterns and support short patterns.'
)
doc.add_paragraph(
    'This means the model adapts in two ways simultaneously: (1) calendar-based features '
    'encode historical tendencies for this time period and election cycle, and (2) real-time '
    'market features capture whether bearish conditions are actually present right now. A '
    'midterm summer where the market is surprisingly strong will score better than a midterm '
    'summer where the expected weakness has materialized, because the real-time features '
    'override the calendar expectation.'
)

doc.add_heading('VIX Hurricane Filter', level=2)
doc.add_paragraph(
    'When VIX exceeds 35, the scoring service refuses to score any opportunity. During '
    'market panics, seasonal patterns break down regardless of quality, direction, or '
    'historical depth. Approximately 4.8% of training samples had VIX > 35 at the '
    'pattern entry date and were removed from the training data entirely, so the model '
    'has no learned behavior for panic regimes. This is a hard safety cutoff, not a '
    'model decision.'
)
doc.add_paragraph(
    'The filter applies only at the entry date. Samples where VIX was below 35 at entry '
    'but spiked above 35 during the holding period are fully included in the training '
    'data. The model learned from those outcomes, meaning the predicted returns and win '
    'probabilities already account for the possibility of a mid-trade VIX spike. The '
    'filter prevents entering new positions during a panic; it does not address what '
    'happens to open positions if a panic starts after entry. That is a position '
    'monitoring concern that belongs in the trade execution layer.'
)

# ============================================================
# 9. BOTTOM LINE
# ============================================================
doc.add_heading('9. Bottom Line', level=1)

p = doc.add_paragraph()
p.add_run('The model provides a consistent filtering edge, not a crystal ball. ').bold = True
p.add_run(
    'It takes seasonal patterns that already have a ~65% base win rate and reliably '
    'pushes them to ~78-82% by identifying which patterns are most likely to work in '
    'current market conditions. The edge is small per trade but consistent across 8 '
    'out-of-sample years spanning very different market regimes.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Strongest endorsement: ').bold = True
p.add_run(
    'The ML_70 cohort (top 30% of scored opportunities, equal-weighted) has never had a '
    'negative average return in any of the 8 walk-forward validation years. This is a '
    'cohort statistic across hundreds of thousands of opportunities per year, not a '
    'portfolio backtest -- but it demonstrates consistent filtering ability across '
    'very different market regimes.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Biggest risk: ').bold = True
p.add_run(
    'Individual trade losses can be large (especially with options), and the model '
    'cannot prevent them. Risk management is not optional -- it IS the strategy. The '
    'model selects; risk management protects.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Production readiness: ').bold = True
p.add_run(
    'Ready for production deployment in TradeWave UI. Ready for controlled live testing '
    'in automated trading (shadow mode or small capital). Full production trading requires '
    'validated trade selection policy, portfolio construction rules, and live monitoring.'
)

# ============================================================
# SAVE
# ============================================================
out_path = os.path.join(os.path.dirname(__file__), 'ML_Scorer_V2_Analysis.docx')
doc.save(out_path)
print(f"Saved to {out_path}")
